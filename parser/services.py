import os
import uuid
from datetime import datetime, timedelta
from django.conf import settings
from supabase import create_client, Client
import logging
from PIL import Image
import io
import fitz  # PyMuPDF for PDF processing
import json
import re
import time
from typing import Dict, Any, Optional, List
import google.generativeai as genai
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError

# File generation imports
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from docx import Document

logger = logging.getLogger(__name__)


class ErrorHandler:
    """Simplified error handling utility"""
    
    @staticmethod
    def success(message, data=None):
        """Format a success response"""
        response = {'success': True, 'message': message}
        if data is not None:
            response['data'] = data
        return response
    
    @staticmethod
    def error(message, retry_allowed=True):
        """Format an error response"""
        return {
            'success': False,
            'error': message,
            'retry_allowed': retry_allowed
        }
    
    @staticmethod
    def get_user_friendly_error(exception):
        """Convert technical exceptions to user-friendly error messages"""
        return {
            'success': False,
            'error': 'Processing failed',
            'retry_allowed': True
        }


class SupabaseStorageService:
    """Service for handling file uploads to Supabase Storage"""
    
    def __init__(self):
        try:
            # Use service role key for server-side operations
            api_key = getattr(settings, 'SUPABASE_SERVICE_KEY', None) or settings.SUPABASE_KEY
            self.supabase: Client = create_client(
                settings.SUPABASE_URL,
                api_key
            )
            self.bucket_name = settings.SUPABASE_BUCKET_NAME
        except Exception:
            self.supabase = None
    
    def upload_file(self, file, session_key):
        """Upload file to Supabase Storage"""
        if not self.supabase:
            return ErrorHandler.error('Storage unavailable')
        
        if not file or file.size > 10 * 1024 * 1024:
            return ErrorHandler.error('Invalid file')
        
        # Generate filename and read content
        unique_filename = f"{session_key}/{uuid.uuid4()}{os.path.splitext(file.name)[1] or '.tmp'}"
        
        try:
            file.seek(0)
            file_content = file.read()
            if not file_content:
                return ErrorHandler.error('Empty file')
        except Exception:
            return ErrorHandler.error('File read failed')
        
        # Upload with 2-attempt retry
        for attempt in range(2):
            try:
                response = self.supabase.storage.from_(self.bucket_name).upload(
                    path=unique_filename,
                    file=file_content,
                    file_options={"content-type": file.content_type or 'application/octet-stream'}
                )
                
                # Check if upload was successful (newer Supabase client doesn't use status_code)
                if response and not hasattr(response, 'error'):
                    try:
                        public_url = self.supabase.storage.from_(self.bucket_name).get_public_url(unique_filename)
                    except Exception:
                        public_url = None
                    
                    return {
                        'success': True,
                        'file_path': unique_filename,
                        'public_url': public_url
                    }
                    
            except Exception as e:
                logger.error(f"Upload attempt {attempt + 1} failed: {str(e)}")
                pass
            
            if attempt == 0:
                time.sleep(1)
        
        return ErrorHandler.error('Upload failed')
    
    def delete_file(self, file_path):
        """Delete file from storage"""
        try:
            response = self.supabase.storage.from_(self.bucket_name).remove([file_path])
            # Newer Supabase client - check if response exists and no error
            return response and not hasattr(response, 'error')
        except Exception:
            return False
    
    def cleanup_session_files(self, session_key):
        """Clean up session files"""
        try:
            files = self.supabase.storage.from_(self.bucket_name).list(session_key)
            if files:
                file_paths = [f"{session_key}/{f['name']}" for f in files]
                response = self.supabase.storage.from_(self.bucket_name).remove(file_paths)
                return response and not hasattr(response, 'error')
            return True
        except Exception:
            return False
    
    def get_file_content(self, file_path):
        """Download file content"""
        try:
            return self.supabase.storage.from_(self.bucket_name).download(file_path)
        except Exception:
            return None


class FileCleanupService:
    """Service for managing file cleanup and storage maintenance"""
    
    def __init__(self):
        self.storage_service = SupabaseStorageService()
    
    def cleanup_expired_files(self, hours_old=1):
        """Clean up files older than specified hours"""
        from .models import ProcessedDocument, UserSession
        from django.utils import timezone
        
        try:
            cutoff_time = timezone.now() - timedelta(hours=hours_old)
            
            # Find and delete old sessions and documents
            old_sessions = UserSession.objects.filter(last_activity__lt=cutoff_time)
            sessions_count = old_sessions.count()
            documents_count = ProcessedDocument.objects.filter(session__in=old_sessions).count()
            
            # Delete old documents and sessions
            ProcessedDocument.objects.filter(session__in=old_sessions).delete()
            old_sessions.delete()
            
            return ErrorHandler.success(f"Cleaned up {sessions_count} sessions")
            
        except Exception as e:
            logger.error(f"Cleanup failed: {str(e)}")
            return ErrorHandler.error(f"Cleanup failed: {str(e)}")
    
    def cleanup_session_manually(self, session_key):
        """Manually clean up a specific session"""
        from .models import ProcessedDocument, UserSession
        
        try:
            # Clean up storage files
            self.storage_service.cleanup_session_files(session_key)
            
            # Clean up database records
            try:
                session = UserSession.objects.get(session_key=session_key)
                documents_count = session.documents.count()
                session.documents.all().delete()
                session.delete()
                return ErrorHandler.success(f"Cleaned up session with {documents_count} documents")
            except UserSession.DoesNotExist:
                return ErrorHandler.success("Session not found")
                
        except Exception as e:
            logger.error(f"Manual cleanup failed: {str(e)}")
            return ErrorHandler.error(f"Manual cleanup failed: {str(e)}")
    
    def schedule_automatic_cleanup(self):
        """Perform automatic cleanup for scheduled maintenance"""
        result = self.cleanup_expired_files(hours_old=1)
        
        if not result.get('success'):
            logger.error(f"Automatic cleanup failed: {result.get('error')}")
        
        return result


class SessionService:
    """Service for managing user sessions and concurrent limits"""
    
    @staticmethod
    def get_or_create_session(request):
        """Get or create user session"""
        from .models import UserSession
        
        session_key = request.session.session_key
        if not session_key:
            request.session.create()
            session_key = request.session.session_key
        
        try:
            user_session = UserSession.objects.get(session_key=session_key)
            if not user_session.is_active:
                user_session.is_active = True
                user_session.save()
            return user_session, False, None
            
        except UserSession.DoesNotExist:
            user_session = UserSession.objects.create(
                session_key=session_key,
                is_active=True
            )
            return user_session, True, None
    
    @staticmethod
    def cleanup_inactive_sessions():
        """Clean up sessions inactive for more than 1 hour"""
        from .models import UserSession
        
        cutoff_time = datetime.now() - timedelta(hours=1)
        inactive_sessions = UserSession.objects.filter(
            last_activity__lt=cutoff_time,
            is_active=True
        )
        
        cleanup_service = FileCleanupService()
        for session in inactive_sessions:
            cleanup_service.cleanup_session_manually(session.session_key)
        
        inactive_sessions.update(is_active=False)


class LLMService:
    """Service for parsing extracted text using Google Gemini API with fallback parsing"""
    
    def __init__(self):
        # Initialize Gemini client
        self.gemini_client = None
        self.vision_client = None
        if hasattr(settings, 'GEMINI_API_KEY') and settings.GEMINI_API_KEY:
            try:
                genai.configure(api_key=settings.GEMINI_API_KEY)
                # Default to Gemini 2.5 Flash, with Lite used elsewhere as fallback
                model_name = getattr(settings, 'GEMINI_MODEL', 'gemini-2.5-flash')
                self.gemini_client = genai.GenerativeModel(model_name)
                vision_model = getattr(settings, 'GEMINI_VISION_MODEL', 'gemini-2.5-flash')
                self.vision_client = genai.GenerativeModel(vision_model)
                logger.info(f"Gemini API initialized successfully with model: {model_name}")
            except Exception as e:
                logger.warning(f"Failed to initialize Gemini API: {str(e)}")
                self.gemini_client = None
                self.vision_client = None
    
    def parse_banking_document(self, text: str, document_type: str = "document", ocr_engine: str = "vision") -> Dict[str, Any]:
        """Parse document text using Gemini API or fallback pattern matching"""
        if not text or not text.strip():
            return {'success': False, 'error': 'No text provided for parsing', 'data': {}}
        
        # Try Gemini API first if available
        if self.gemini_client:
            logger.info(f"Using Gemini API for document parsing (engine: {ocr_engine})")
            gemini_result = self._try_gemini_parsing(text, document_type, ocr_engine)
            if gemini_result['success']:
                return gemini_result
            else:
                logger.warning(f"Gemini API failed: {gemini_result.get('error')}. Falling back to pattern matching.")
        
        # Fallback to pattern matching
        logger.info("Using pattern matching fallback for document parsing")
        return self._fallback_pattern_parsing(text)
    
    def process_document_with_vision(self, image_data: bytes, file_type: str) -> Dict[str, Any]:
        """Process document directly with Gemini Vision API"""
        if not self.vision_client:
            return {'success': False, 'error': 'Gemini Vision API not available', 'data': {}}
        
        try:
            # Convert image data to PIL Image
            image = Image.open(io.BytesIO(image_data))
            
            # Build vision prompt
            prompt = self._build_vision_prompt()
            
            # Process with Gemini Vision
            response = self.vision_client.generate_content([prompt, image])

            # Safely extract text from Gemini response
            resp_text = ""
            try:
                resp_text = (getattr(response, "text", None) or "").strip()
            except Exception:
                resp_text = ""
            if not resp_text:
                try:
                    for cand in getattr(response, "candidates", []) or []:
                        content = getattr(cand, "content", None)
                        parts = getattr(content, "parts", None) if content else None
                        if not parts:
                            continue
                        buf = []
                        for p in parts:
                            t = getattr(p, "text", None)
                            if t:
                                buf.append(t)
                            elif isinstance(p, dict) and p.get("text"):
                                buf.append(str(p["text"]))
                        if buf:
                            resp_text = "\n".join(buf)
                            break
                except Exception:
                    resp_text = ""

            if not resp_text:
                return {'success': False, 'error': 'No response from Gemini Vision', 'data': {}}
            
            # Parse the response
            parsed_data = self._parse_llm_response(resp_text)
            if parsed_data:
                return {
                    'success': True,
                    'data': parsed_data,
                    'message': 'Document processed successfully with Gemini Vision',
                    'raw_text': response.text[:500] + '...' if len(response.text) > 500 else response.text
                }
            else:
                return {'success': False, 'error': 'Failed to parse Gemini Vision response', 'data': {}}
                
        except Exception as e:
            logger.error(f"Gemini Vision processing failed: {str(e)}")
            return {'success': False, 'error': f'Vision processing failed: {str(e)}', 'data': {}}
    
    def _try_gemini_parsing(self, text: str, document_type: str, ocr_engine: str = "tesseract") -> Dict[str, Any]:
        """Try parsing with Gemini API"""
        # Simple retry logic - try twice
        for attempt in range(2):
            try:
                prompt = self._build_parsing_prompt(text, document_type, ocr_engine)
                
                # Run the API call with a timeout to prevent hangs
                def _call_model():
                    return self.gemini_client.generate_content(prompt)
                
                with ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(_call_model)
                    response = future.result(timeout=90)  # Increased timeout for better models
                
                # Safely extract text; skip if none
                resp_text = ""
                try:
                    resp_text = (getattr(response, "text", None) or "").strip()
                except Exception:
                    resp_text = ""
                if not resp_text:
                    try:
                        for cand in getattr(response, "candidates", []) or []:
                            content = getattr(cand, "content", None)
                            parts = getattr(content, "parts", None) if content else None
                            if not parts:
                                continue
                            buf = []
                            for p in parts:
                                t = getattr(p, "text", None)
                                if t:
                                    buf.append(t)
                                elif isinstance(p, dict) and p.get("text"):
                                    buf.append(str(p["text"]))
                            if buf:
                                resp_text = "\n".join(buf)
                                break
                    except Exception:
                        resp_text = ""
                if not resp_text:
                    continue
                
                parsed_data = self._parse_llm_response(resp_text)
                if parsed_data:
                    return {
                        'success': True,
                        'data': parsed_data,
                        'message': f'Document parsed successfully with Gemini API (OCR: {ocr_engine})'
                    }
                    
            except FuturesTimeoutError:
                logger.error(f"Gemini API timeout (attempt {attempt + 1})")
                if attempt == 1:
                    return {'success': False, 'error': 'AI service timeout', 'data': {}}
            except Exception as e:
                logger.error(f"Gemini API error (attempt {attempt + 1}): {str(e)}")
                if attempt == 1:  # Last attempt
                    return {'success': False, 'error': f'AI service error: {str(e)}', 'data': {}}
        
        return {'success': False, 'error': 'Failed to parse document with Gemini', 'data': {}}
    
    def _fallback_pattern_parsing(self, text: str) -> Dict[str, Any]:
        """Fallback parsing using pattern matching for common document fields"""
        try:
            # Clean the text
            text_lines = [line.strip() for line in text.split('\n') if line.strip()]
            full_text = ' '.join(text_lines)
            
            # Extract basic information using patterns
            parsed_data = {
                'document_type': self._detect_document_type(full_text),
                'confidence_score': 0.7,  # Lower confidence for pattern matching
                'personal_information': self._extract_personal_info(full_text, text_lines),
                'financial_data': self._extract_financial_data(full_text, text_lines),
                'dates': self._extract_dates(full_text),
                'bank_information': self._extract_bank_info(full_text)
            }
            
            return {
                'success': True,
                'data': parsed_data,
                'message': 'Document parsed successfully with pattern matching'
            }
            
        except Exception as e:
            logger.error(f"Fallback parsing failed: {str(e)}")
            return {
                'success': False,
                'error': f'Pattern parsing failed: {str(e)}',
                'data': {}
            }
    
    def _detect_document_type(self, text: str) -> str:
        """Detect document type from text content"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['statement', 'bank statement', 'account statement']):
            return 'Bank Statement'
        elif any(word in text_lower for word in ['balance', 'account balance']):
            return 'Balance Inquiry'
        elif any(word in text_lower for word in ['loan', 'credit', 'mortgage']):
            return 'Loan Document'
        elif any(word in text_lower for word in ['transaction', 'payment', 'transfer']):
            return 'Transaction Record'
        else:
            return 'Document'
    
    def _extract_personal_info(self, full_text: str, text_lines: List[str]) -> Dict[str, Any]:
        """Extract personal information using patterns"""
        personal_info = {}
        
        # Look for account numbers (8-16 digits, sometimes with spaces or dashes)
        account_patterns = [
            r'account[:\s]*([0-9\s\-]{8,20})',
            r'a/c[:\s]*([0-9\s\-]{8,20})',
            r'acc[:\s]*([0-9\s\-]{8,20})'
        ]
        
        for pattern in account_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                account_num = re.sub(r'[^\d]', '', match.group(1))
                if len(account_num) >= 8:
                    personal_info['account_number'] = account_num
                    break
        
        # Look for names (capitalized words, usually at the beginning)
        name_patterns = [
            r'name[:\s]*([A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)*)',
            r'customer[:\s]*([A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)*)'
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                personal_info['full_name'] = match.group(1).strip()
                break
        
        # If no pattern match, try to find capitalized names in first few lines
        if 'full_name' not in personal_info:
            for line in text_lines[:5]:
                # Look for lines with 2+ capitalized words
                words = line.split()
                cap_words = [w for w in words if w and w[0].isupper() and len(w) > 1]
                if len(cap_words) >= 2 and len(' '.join(cap_words)) < 50:
                    personal_info['full_name'] = ' '.join(cap_words)
                    break
        
        return personal_info
    
    def _extract_financial_data(self, full_text: str, text_lines: List[str]) -> Dict[str, Any]:
        """Extract financial information"""
        financial_data = {}
        transactions = []
        
        # Look for balance amounts
        balance_patterns = [
            r'balance[:\s]*([A-Z]{0,3})\s*([0-9,]+\.?\d*)',
            r'available[:\s]*([A-Z]{0,3})\s*([0-9,]+\.?\d*)',
            r'current[:\s]*([A-Z]{0,3})\s*([0-9,]+\.?\d*)'
        ]
        
        for pattern in balance_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                currency = match.group(1) if match.group(1) else 'ETB'
                amount = match.group(2).replace(',', '')
                financial_data['account_balance'] = f"{currency} {amount}"
                break
        
        # Look for transactions in table-like format
        for line in text_lines:
            # Pattern for transaction-like lines (date, description, amount)
            trans_match = re.search(r'(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})\s+(.+?)\s+([+-]?\d+[,.]?\d*)', line)
            if trans_match:
                date_str = trans_match.group(1)
                description = trans_match.group(2).strip()
                amount = trans_match.group(3).replace(',', '')
                
                transaction_type = 'debit' if amount.startswith('-') or float(amount.replace('-', '')) < 0 else 'credit'
                
                transactions.append({
                    'date': date_str,
                    'description': description,
                    'amount': amount,
                    'type': transaction_type
                })
        
        if transactions:
            financial_data['transactions'] = transactions
        
        return financial_data
    
    def _extract_dates(self, text: str) -> Dict[str, Any]:
        """Extract important dates"""
        dates = {}
        
        # Look for statement dates
        date_patterns = [
            r'statement\s+date[:\s]*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'date[:\s]*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                dates['statement_date'] = match.group(1)
                break
        
        return dates
    
    def _extract_bank_info(self, text: str) -> Dict[str, Any]:
        """Extract bank information"""
        bank_info = {}
        
        # Common Ethiopian banks
        ethiopian_banks = [
            'Commercial Bank of Ethiopia', 'CBE', 'Dashen Bank', 'Awash Bank',
            'Bank of Abyssinia', 'Wegagen Bank', 'United Bank', 'Nib Bank',
            'Cooperative Bank of Oromia', 'Lion Bank', 'Oromia Bank',
            'Abay Bank', 'Addis International Bank', 'Debub Global Bank'
        ]
        
        text_upper = text.upper()
        for bank in ethiopian_banks:
            if bank.upper() in text_upper:
                bank_info['bank_name'] = bank
                break
        
        # If no specific bank found, look for "bank" keyword
        if 'bank_name' not in bank_info:
            bank_match = re.search(r'([A-Z][a-z]+\s+Bank)', text)
            if bank_match:
                bank_info['bank_name'] = bank_match.group(1)
        
        return bank_info
    

    
    def _build_parsing_prompt(self, text: str, document_type: str, ocr_engine: str = "vision") -> str:
        """Prompt for structuring text while preserving exact character representation."""
        return f"""You are an expert multilingual financial document parser.

Task: Structure the provided text into JSON format WITHOUT modifying any characters, words, or formatting. If no table can be inferred, fall back to a single table named 'main' with headers ['text'] and rows = each input line as a separate row (preserve order).

CRITICAL CHARACTER PRESERVATION RULES:
- NEVER autocorrect, fix, or modify any characters, words, or text from the input
- NEVER fix what appears to be OCR errors or typos - preserve them EXACTLY
- NEVER transliterate Amharic/Ethiopic characters to Latin script
- NEVER normalize or standardize formatting - keep original spacing, punctuation
- NEVER correct obvious mistakes like 0/O, 1/l, 5/S - transcribe exactly as provided
- NEVER standardize dates or numbers - keep original format exactly
- NEVER add missing punctuation or correct grammar
- NEVER change case (uppercase/lowercase) from what is provided

STRUCTURING RULES:
- Copy each character, symbol, and space EXACTLY as provided in the input text
- For tables: preserve headers exactly as written; maintain original column structure
- For key:value pairs: output as two-column table [key, value] with exact text
- For lists/paragraphs: output as single-column table preserving original line breaks
- Maintain original text order and formatting
- If structure is unclear, default to single-column table with original text

Output STRICT JSON only with this schema:
{{
  "tables": [ {{ "name": string, "headers": [string], "rows": [[string]] }} ]
}}

Document type: {document_type}
Input text to preserve exactly:
{text}
"""
    
    def _build_vision_prompt(self) -> str:
        """Build character-preserving prompt for Gemini Vision API"""
        return """Extract all text from this document image with ABSOLUTE CHARACTER PRESERVATION.

If you cannot confidently detect a table, produce a single table named 'main' with headers ['text'] and rows = each visual line as a separate row (preserve order).

CRITICAL CHARACTER PRESERVATION RULES:
- NEVER autocorrect, fix, or modify any characters, words, or text you see
- NEVER fix what appears to be OCR errors, typos, or misspellings - transcribe EXACTLY
- NEVER transliterate Amharic/Ethiopian characters to Latin script - preserve original script
- NEVER normalize or standardize formatting - keep original spacing and punctuation
- NEVER correct obvious mistakes like 0/O, 1/l, 5/S - transcribe exactly as shown in image
- NEVER standardize dates, numbers, or currency - keep original format (e.g., 12/5/23, not 2023-05-12)
- NEVER add missing punctuation or correct grammar
- NEVER change case (uppercase/lowercase) from what is visible
- NEVER interpret or translate abbreviations - copy exactly as shown

EXTRACTION INSTRUCTIONS:
- Copy each visible character, symbol, digit, and space EXACTLY as it appears
- Preserve all original formatting, spacing, line breaks, and alignment
- Include all visible text: printed, handwritten, stamps, signatures, watermarks
- Maintain exact table structure with original headers and cell content
- Preserve spatial relationships and reading order (top to bottom, left to right)
- If text is unclear, transcribe your best visual interpretation without correcting

Return structured JSON that contains the document's exact visual content without any modifications."""
    
    def _parse_llm_response(self, response_text: str) -> Optional[Dict[str, Any]]:
        """Parse LLM response and extract JSON data - flexible parsing"""
        try:
            cleaned_text = response_text.strip()
            
            # Try to find JSON in the response
            json_start = cleaned_text.find('{')
            json_end = cleaned_text.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_text = cleaned_text[json_start:json_end]
                parsed_data = json.loads(json_text)
                
                # If parsed successfully, return as-is (no template enforcement)
                return parsed_data
            
            # If no JSON found, create a simple structure from the text
            return {
                "document_type": "Document",
                "extracted_text": cleaned_text[:1000],  # First 1000 chars
                "parsing_method": "text_fallback"
            }
                
        except json.JSONDecodeError:
            # Fallback: return the text as extracted content
            return {
                "document_type": "Document", 
                "extracted_text": response_text[:1000],
                "parsing_method": "text_fallback"
            }
        except Exception:
            return None
    
    def test_api_connection(self) -> Dict[str, Any]:
        """Test Gemini API connection"""
        if not self.gemini_client:
            return {'gemini': {'available': False, 'error': 'API key not configured'}}
        
        try:
            test_response = self.gemini_client.generate_content("Test connection. Respond with 'OK'.")
            if test_response.text and 'OK' in test_response.text:
                return {'gemini': {'available': True, 'error': None}}
            else:
                return {'gemini': {'available': False, 'error': 'Unexpected response'}}
        except Exception as e:
            return {'gemini': {'available': False, 'error': str(e)}}


class DataStructuringService:
    """Service for organizing and formatting extracted document data"""
    
    def structure_document_data(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Structure parsed document data exactly as extracted - no metadata added"""
        try:
            # Return the parsed data as-is, without adding metadata or processing timestamps
            structured_data = {}
            
            # Dynamically organize whatever data was extracted
            for key, value in parsed_data.items():
                if key in ['document_type', 'parsing_method']:
                    continue  # Skip internal processing fields
                
                if isinstance(value, dict) and value:
                    structured_data[key] = self._format_data(value)
                elif isinstance(value, list) and value:
                    # Handle any list data - could be transactions, items, entries, etc.
                    if self._looks_like_tabular_data(value):
                        structured_data[key] = self._format_tabular_data(value)
                    else:
                        structured_data[key] = value
                elif value:  # Any other non-empty value
                    structured_data[key] = [{'field': key.replace('_', ' ').title(), 'value': str(value)}]
            
            return ErrorHandler.success('Data structured successfully', structured_data)
            
        except Exception as e:
            return ErrorHandler.error(f'Data structuring failed: {str(e)}')
    
    def _format_data(self, data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Format data section as key-value pairs"""
        def _normalize_value(value: Any) -> str:
            if value is None:
                return 'N/A'
            if isinstance(value, str):
                return value.strip() or 'N/A'
            return str(value)
        return [
            {'field': key.replace('_', ' ').title(), 'value': _normalize_value(value)}
            for key, value in data.items()
        ]
    
    def _looks_like_tabular_data(self, data: List[Any]) -> bool:
        """Check if list data looks like tabular data (dictionaries with similar keys)"""
        if not data or len(data) < 2:
            return False
        
        # Check if first few items are dictionaries with similar structure
        sample_items = data[:3]
        if not all(isinstance(item, dict) for item in sample_items):
            return False
        
        # Check for common keys that suggest tabular data
        first_keys = set(sample_items[0].keys())
        return len(first_keys) > 1 and all(
            len(set(item.keys()) & first_keys) / len(first_keys) > 0.5 
            for item in sample_items[1:]
        )
    
    def _format_tabular_data(self, data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Format tabular data (transactions, entries, records, etc.)"""
        return [
            {
                'id': i + 1,
                **{k: str(v) if v is not None else '' for k, v in item.items()}
            }
            for i, item in enumerate(data[:100])  # Limit to 100 entries
        ]



class FileGenerationService:
    """Simplified service for generating output files in Excel, PDF, and DOC formats"""
    
    def __init__(self):
        self.temp_dir = os.path.join(settings.BASE_DIR, 'temp_files')
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def generate_all_formats(self, structured_data: Dict[str, Any], session_key: str) -> Dict[str, Any]:
        """Generate all three output formats with simplified error handling"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"document_{session_key}_{timestamp}"
        
        results = {'success': True, 'files': {}, 'errors': []}
        
        # Generate files with basic error handling
        for format_type, generator_method in [
            ('excel', self.generate_excel_file),
            ('pdf', self.generate_pdf_file), 
            ('doc', self.generate_doc_file)
        ]:
            result = generator_method(structured_data, base_filename)
            if result.get('success'):
                results['files'][format_type] = {
                    'path': result['path'],
                    'filename': result['filename'],
                    'size': os.path.getsize(result['path']) if os.path.exists(result['path']) else 0
                }
            else:
                results['errors'].append(f"{format_type.upper()} generation failed")
                results['success'] = False
        
        return results
    
    def generate_excel_file(self, data: Dict[str, Any], base_filename: str) -> Dict[str, Any]:
        """Generate Excel with proper tables, key-value headers, and widths."""
        try:
            wb = Workbook()
            # remove default sheet; we'll add sheets as needed
            default_ws = wb.active
            wb.remove(default_ws)

            def autosize(ws):
                for col_idx in range(1, ws.max_column + 1):
                    letter = get_column_letter(col_idx)
                    max_len = 0
                    for cell in ws.iter_rows(min_col=col_idx, max_col=col_idx, values_only=True):
                        val = cell[0]
                        if val is None:
                            continue
                        s = str(val)
                        if len(s) > max_len:
                            max_len = len(s)
                    ws.column_dimensions[letter].width = min(max(10, max_len + 2), 60)

            # 1) Direct support for LLM "tables" schema
            if isinstance(data, dict) and isinstance(data.get('tables'), list) and data['tables']:
                for tbl in data['tables']:
                    title = (tbl.get('name') or 'Sheet')[:31]
                    ws = wb.create_sheet(title)
                    headers = tbl.get('headers') or []
                    rows = tbl.get('rows') or []
                    if headers:
                        ws.append(headers)
                        for i in range(1, len(headers) + 1):
                            ws.cell(row=1, column=i).font = Font(bold=True)
                        ws.freeze_panes = 'A2'
                    for r in rows:
                        ws.append(["" if v is None else v for v in r])
                    autosize(ws)

            # 2) Fallback: iterate other sections
            else:
                ws = wb.create_sheet('Data')
                row = 1
                for section_name, section_data in data.items():
                    if not section_data or section_name == 'metadata':
                        continue
                    # tabular list of dicts (with or without id)
                    if isinstance(section_data, list) and section_data and isinstance(section_data[0], dict):
                        headers = [k for k in section_data[0].keys() if k != 'id'] or list(section_data[0].keys())
                        ws.append(headers)
                        for i in range(1, len(headers) + 1):
                            ws.cell(row=row, column=i).font = Font(bold=True)
                        ws.freeze_panes = f'A{row+1}'
                        row += 1
                        for item in section_data:
                            ws.append([str(item.get(h, '')) for h in headers])
                            row += 1
                    elif isinstance(section_data, list):
                        # list of primitives or dicts with field/value
                        # ensure headers Field/Value when dicts present
                        has_kv = any(isinstance(x, dict) and ('field' in x or 'value' in x) for x in section_data)
                        if has_kv:
                            ws.append(['Field', 'Value'])
                            ws.cell(row=row, column=1).font = Font(bold=True)
                            ws.cell(row=row, column=2).font = Font(bold=True)
                            ws.freeze_panes = f'A{row+1}'
                            row += 1
                            for x in section_data:
                                if isinstance(x, dict):
                                    ws.append([str(x.get('field', '')), str(x.get('value', ''))])
                                    row += 1
                        else:
                            for x in section_data:
                                ws.cell(row=row, column=1, value=str(x))
                                row += 1
                    elif isinstance(section_data, dict):
                        # two-column key:value
                        ws.append(['Field', 'Value'])
                        ws.cell(row=row, column=1).font = Font(bold=True)
                        ws.cell(row=row, column=2).font = Font(bold=True)
                        ws.freeze_panes = f'A{row+1}'
                        row += 1
                        for k, v in section_data.items():
                            ws.append([str(k).title().replace('_', ' '), '' if v is None else str(v)])
                            row += 1
                autosize(ws)

            output_path = os.path.join(self.temp_dir, f"{base_filename}.xlsx")
            wb.save(output_path)
            return {'success': True, 'path': output_path, 'filename': os.path.basename(output_path)}
        except Exception as e:
            return ErrorHandler.error(f'Excel generation failed: {str(e)}')
    
    def generate_pdf_file(self, data: Dict[str, Any], base_filename: str) -> Dict[str, Any]:
        """Generate PDF with headers and proper tables when available."""
        try:
            output_path = os.path.join(self.temp_dir, f"{base_filename}.pdf")
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            def table_with_header(headers, rows):
                tbl = Table([headers] + rows)
                tbl.setStyle(TableStyle([
                    ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                    ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ]))
                return tbl

            # 1) Direct support for LLM "tables" schema
            if isinstance(data, dict) and isinstance(data.get('tables'), list) and data['tables']:
                for tbl in data['tables']:
                    headers = tbl.get('headers') or []
                    rows = [["" if v is None else str(v) for v in r] for r in (tbl.get('rows') or [])]
                    if headers or rows:
                        story.append(table_with_header(headers or [''], rows or [['']]))
                        story.append(Spacer(1, 12))
            else:
                # Fallback: render sections
                for section_name, section_data in data.items():
                    if not section_data or section_name == 'metadata':
                        continue
                    if isinstance(section_data, list) and section_data and isinstance(section_data[0], dict):
                        headers = [k for k in section_data[0].keys() if k != 'id'] or list(section_data[0].keys())
                        rows = [[str(item.get(h, '')) for h in headers] for item in section_data]
                        story.append(table_with_header(headers, rows))
                        story.append(Spacer(1, 12))
                    elif isinstance(section_data, list):
                        has_kv = any(isinstance(x, dict) and ('field' in x or 'value' in x) for x in section_data)
                        if has_kv:
                            headers = ['Field', 'Value']
                            rows = [[str(x.get('field', '')), str(x.get('value', ''))] for x in section_data if isinstance(x, dict)]
                            story.append(table_with_header(headers, rows))
                            story.append(Spacer(1, 12))
                        else:
                            for x in section_data:
                                story.append(Paragraph(str(x), styles['Normal']))
                    elif isinstance(section_data, dict):
                        headers = ['Field', 'Value']
                        rows = [[str(k).title().replace('_', ' '), '' if v is None else str(v)] for k, v in section_data.items()]
                        story.append(table_with_header(headers, rows))
                        story.append(Spacer(1, 12))

            doc.build(story)
            return {'success': True, 'path': output_path, 'filename': os.path.basename(output_path)}
        except Exception as e:
            return ErrorHandler.error(f'PDF generation failed: {str(e)}')
    
    def generate_doc_file(self, data: Dict[str, Any], base_filename: str) -> Dict[str, Any]:
        """Generate DOC file with raw data only"""
        try:
            output_path = os.path.join(self.temp_dir, f"{base_filename}.docx")
            doc = Document()
            
            # Add sections without titles - just raw data
            for section_name, section_data in data.items():
                if not section_data or section_name == 'metadata':
                    continue
                
                if isinstance(section_data, list):
                    # Check if this is tabular data
                    if section_data and isinstance(section_data[0], dict) and 'id' in section_data[0]:
                        # Create table with just data, no headers
                        first_item = section_data[0]
                        headers = [key for key in first_item.keys() if key != 'id']
                        
                        table = doc.add_table(rows=0, cols=len(headers))
                        
                        for item in section_data:
                            if isinstance(item, dict):
                                row_cells = table.add_row().cells
                                for i, header in enumerate(headers):
                                    value = item.get(header, '')
                                    row_cells[i].text = str(value) if value is not None else ''
                    else:
                        # Just raw values without field names
                        for item in section_data:
                            if isinstance(item, dict):
                                value = item.get('value', '')
                                if value:
                                    doc.add_paragraph(str(value))
                elif isinstance(section_data, dict):
                    # Just raw values
                    for key, value in section_data.items():
                        if value:
                            doc.add_paragraph(str(value))
            
            doc.save(output_path)
            
            return {
                'success': True,
                'path': output_path,
                'filename': os.path.basename(output_path)
            }
        except Exception as e:
            return ErrorHandler.error(f'DOC generation failed: {str(e)}')
    
    def cleanup_temp_files(self, file_paths: List[str]):
        """Clean up temporary files"""
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception:
                pass

