import os
import uuid
from datetime import datetime, timedelta
from django.conf import settings
from supabase import create_client, Client
import logging
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import io
import fitz  # PyMuPDF for PDF processing
import json
import re
import time
from typing import Dict, Any, Optional, List
import openai
import google.generativeai as genai
import requests

# File generation imports
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


class ErrorHandler:
    """Centralized error handling and user feedback utility"""
    
    @staticmethod
    def format_error_response(error_type: str, message: str, details: str = None, 
                            suggestions: List[str] = None, retry_allowed: bool = True,
                            fallback_suggestion: str = None) -> Dict[str, Any]:
        """
        Format a standardized error response
        
        Args:
            error_type: Type of error (user-friendly)
            message: Main error message
            details: Technical details about the error
            suggestions: List of actionable suggestions for the user
            retry_allowed: Whether the user should be allowed to retry
            fallback_suggestion: Alternative approach if main method fails
            
        Returns:
            dict: Standardized error response
        """
        response = {
            'success': False,
            'error': error_type,
            'message': message,
            'retry_allowed': retry_allowed,
            'timestamp': datetime.now().isoformat()
        }
        
        if details:
            response['details'] = details
        
        if suggestions:
            response['suggestions'] = suggestions if isinstance(suggestions, list) else [suggestions]
        
        if fallback_suggestion:
            response['fallback_suggestion'] = fallback_suggestion
            
        return response
    
    @staticmethod
    def format_success_response(message: str, data: Any = None, **kwargs) -> Dict[str, Any]:
        """
        Format a standardized success response
        
        Args:
            message: Success message
            data: Response data
            **kwargs: Additional response fields
            
        Returns:
            dict: Standardized success response
        """
        response = {
            'success': True,
            'message': message,
            'timestamp': datetime.now().isoformat()
        }
        
        if data is not None:
            response['data'] = data
            
        response.update(kwargs)
        return response
    
    @staticmethod
    def get_user_friendly_error(exception: Exception) -> Dict[str, Any]:
        """
        Convert technical exceptions to user-friendly error messages
        
        Args:
            exception: The exception that occurred
            
        Returns:
            dict: User-friendly error response
        """
        error_msg = str(exception).lower()
        
        # Network-related errors
        if any(term in error_msg for term in ['connection', 'network', 'timeout', 'unreachable']):
            return ErrorHandler.format_error_response(
                'Connection failed',
                'Could not connect to the service',
                str(exception),
                ['Check your internet connection', 'Try again in a few minutes'],
                retry_allowed=True
            )
        
        # Permission/authentication errors
        if any(term in error_msg for term in ['permission', 'unauthorized', 'forbidden', 'access denied']):
            return ErrorHandler.format_error_response(
                'Access denied',
                'Insufficient permissions to perform this operation',
                str(exception),
                ['Contact system administrator', 'Check your account permissions'],
                retry_allowed=False
            )
        
        # File-related errors
        if any(term in error_msg for term in ['file not found', 'no such file', 'cannot open']):
            return ErrorHandler.format_error_response(
                'File not found',
                'The requested file could not be found',
                str(exception),
                ['Check that the file exists', 'Try uploading the file again'],
                retry_allowed=True
            )
        
        # Memory/resource errors
        if any(term in error_msg for term in ['memory', 'out of space', 'disk full']):
            return ErrorHandler.format_error_response(
                'System resources unavailable',
                'Not enough system resources to complete the operation',
                str(exception),
                ['Try processing a smaller file', 'Try again later', 'Contact administrator'],
                retry_allowed=True
            )
        
        # Generic error
        return ErrorHandler.format_error_response(
            'Processing failed',
            'An unexpected error occurred',
            str(exception),
            ['Try again in a few minutes', 'Contact support if problem persists'],
            retry_allowed=True
        )


class SupabaseStorageService:
    """Service for handling file uploads to Supabase Storage"""
    
    def __init__(self):
        try:
            self.supabase: Client = create_client(
                settings.SUPABASE_URL,
                settings.SUPABASE_KEY
            )
            self.bucket_name = "document-uploads"
        except Exception as e:
            logger.error(f"Failed to initialize Supabase client: {str(e)}")
            self.supabase = None
    
    def upload_file(self, file, session_key):
        """
        Upload file to Supabase Storage with comprehensive error handling
        
        Args:
            file: Django UploadedFile object
            session_key: User session key for organizing files
            
        Returns:
            dict: Contains file_path and public_url with detailed error info
        """
        if not self.supabase:
            return {
                'success': False,
                'error': 'Storage service unavailable',
                'details': 'Database connection could not be established',
                'suggestions': ['Check your internet connection', 'Try again in a few moments'],
                'retry_allowed': True
            }
        
        try:
            # Validate file
            if not file or not hasattr(file, 'read'):
                return {
                    'success': False,
                    'error': 'Invalid file provided',
                    'details': 'File object is missing or corrupted',
                    'suggestions': ['Please select a valid file', 'Try uploading a different file'],
                    'retry_allowed': True
                }
            
            # Check file size
            if file.size > 10 * 1024 * 1024:  # 10MB
                return {
                    'success': False,
                    'error': 'File too large',
                    'details': f'File size is {file.size / (1024*1024):.1f}MB, maximum allowed is 10MB',
                    'suggestions': ['Compress your file', 'Use a smaller image resolution', 'Split large documents'],
                    'retry_allowed': False
                }
            
            # Generate unique filename
            file_extension = os.path.splitext(file.name)[1]
            if not file_extension:
                file_extension = '.tmp'
            unique_filename = f"{session_key}/{uuid.uuid4()}{file_extension}"
            
            # Read file content with error handling
            try:
                file.seek(0)  # Reset file pointer
                file_content = file.read()
                if not file_content:
                    return {
                        'success': False,
                        'error': 'Empty file',
                        'details': 'The uploaded file appears to be empty',
                        'suggestions': ['Check that your file contains data', 'Try uploading a different file'],
                        'retry_allowed': True
                    }
            except Exception as read_error:
                return {
                    'success': False,
                    'error': 'File read error',
                    'details': f'Could not read file content: {str(read_error)}',
                    'suggestions': ['The file may be corrupted', 'Try uploading a different file'],
                    'retry_allowed': True
                }
            
            # Upload to Supabase Storage with retry logic
            max_retries = 2
            for attempt in range(max_retries + 1):
                try:
                    response = self.supabase.storage.from_(self.bucket_name).upload(
                        path=unique_filename,
                        file=file_content,
                        file_options={
                            "content-type": file.content_type or 'application/octet-stream',
                            "cache-control": "3600"
                        }
                    )
                    
                    if response.status_code == 200:
                        # Get public URL
                        try:
                            public_url = self.supabase.storage.from_(self.bucket_name).get_public_url(unique_filename)
                            return {
                                'success': True,
                                'file_path': unique_filename,
                                'public_url': public_url,
                                'message': 'File uploaded successfully',
                                'attempt': attempt + 1
                            }
                        except Exception as url_error:
                            logger.warning(f"Could not get public URL: {str(url_error)}")
                            return {
                                'success': True,
                                'file_path': unique_filename,
                                'public_url': None,
                                'message': 'File uploaded successfully (URL generation failed)',
                                'attempt': attempt + 1
                            }
                    else:
                        if attempt < max_retries:
                            logger.warning(f"Upload attempt {attempt + 1} failed, retrying...")
                            time.sleep(1)  # Brief delay before retry
                            continue
                        else:
                            logger.error(f"Supabase upload failed after {max_retries + 1} attempts: {response}")
                            return {
                                'success': False,
                                'error': 'Upload failed after multiple attempts',
                                'details': f'Server responded with status {response.status_code}',
                                'suggestions': ['Check your internet connection', 'Try again in a few minutes', 'Contact support if problem persists'],
                                'retry_allowed': True
                            }
                            
                except requests.exceptions.ConnectionError:
                    if attempt < max_retries:
                        logger.warning(f"Connection error on attempt {attempt + 1}, retrying...")
                        time.sleep(2)  # Longer delay for connection issues
                        continue
                    else:
                        return {
                            'success': False,
                            'error': 'Connection failed',
                            'details': 'Could not connect to storage service',
                            'suggestions': ['Check your internet connection', 'Try again in a few minutes'],
                            'retry_allowed': True
                        }
                except requests.exceptions.Timeout:
                    if attempt < max_retries:
                        logger.warning(f"Timeout on attempt {attempt + 1}, retrying...")
                        time.sleep(2)
                        continue
                    else:
                        return {
                            'success': False,
                            'error': 'Upload timeout',
                            'details': 'The upload took too long to complete',
                            'suggestions': ['Check your internet connection', 'Try uploading a smaller file', 'Try again later'],
                            'retry_allowed': True
                        }
                
        except Exception as e:
            logger.error(f"Unexpected error uploading file to Supabase: {str(e)}")
            error_msg = str(e).lower()
            
            # Provide specific error messages for common issues
            if 'permission' in error_msg or 'unauthorized' in error_msg:
                return {
                    'success': False,
                    'error': 'Storage access denied',
                    'details': 'Insufficient permissions to upload files',
                    'suggestions': ['Contact system administrator', 'Check service configuration'],
                    'retry_allowed': False
                }
            elif 'network' in error_msg or 'connection' in error_msg:
                return {
                    'success': False,
                    'error': 'Network error',
                    'details': 'Could not connect to storage service',
                    'suggestions': ['Check your internet connection', 'Try again in a few minutes'],
                    'retry_allowed': True
                }
            else:
                return {
                    'success': False,
                    'error': 'Storage service error',
                    'details': str(e),
                    'suggestions': ['Try again in a few minutes', 'Contact support if problem persists'],
                    'retry_allowed': True
                }
    
    def delete_file(self, file_path):
        """
        Delete file from Supabase Storage
        
        Args:
            file_path: Path to file in storage
            
        Returns:
            bool: Success status
        """
        try:
            response = self.supabase.storage.from_(self.bucket_name).remove([file_path])
            return response.status_code == 200
        except Exception as e:
            logger.error(f"Error deleting file from Supabase: {str(e)}")
            return False
    
    def cleanup_session_files(self, session_key):
        """
        Clean up all files for a specific session
        
        Args:
            session_key: Session key to clean up
            
        Returns:
            bool: Success status
        """
        try:
            # List all files in the session folder
            response = self.supabase.storage.from_(self.bucket_name).list(session_key)
            
            if response:
                file_paths = [f"{session_key}/{file['name']}" for file in response]
                if file_paths:
                    delete_response = self.supabase.storage.from_(self.bucket_name).remove(file_paths)
                    return delete_response.status_code == 200
            
            return True
        except Exception as e:
            logger.error(f"Error cleaning up session files: {str(e)}")
            return False
    
    def get_file_content(self, file_path):
        """
        Download file content from Supabase Storage
        
        Args:
            file_path: Path to file in storage
            
        Returns:
            bytes: File content or None if error
        """
        try:
            response = self.supabase.storage.from_(self.bucket_name).download(file_path)
            return response
        except Exception as e:
            logger.error(f"Error downloading file from Supabase: {str(e)}")
            return None
    
    def list_all_files(self, prefix=""):
        """
        List all files in storage with optional prefix filter
        
        Args:
            prefix: Optional prefix to filter files (e.g., session_key)
            
        Returns:
            list: List of file objects with metadata
        """
        try:
            if not self.supabase:
                return []
            
            response = self.supabase.storage.from_(self.bucket_name).list(prefix)
            return response if response else []
        except Exception as e:
            logger.error(f"Error listing files from Supabase: {str(e)}")
            return []
    
    def get_file_info(self, file_path):
        """
        Get file information including size and last modified date
        
        Args:
            file_path: Path to file in storage
            
        Returns:
            dict: File information or None if error
        """
        try:
            if not self.supabase:
                return None
            
            # Get file list to find the specific file
            path_parts = file_path.split('/')
            if len(path_parts) > 1:
                prefix = '/'.join(path_parts[:-1])
                filename = path_parts[-1]
            else:
                prefix = ""
                filename = file_path
            
            files = self.list_all_files(prefix)
            for file_info in files:
                if file_info.get('name') == filename:
                    return {
                        'name': file_info.get('name'),
                        'size': file_info.get('metadata', {}).get('size', 0),
                        'last_modified': file_info.get('updated_at'),
                        'created_at': file_info.get('created_at')
                    }
            return None
        except Exception as e:
            logger.error(f"Error getting file info from Supabase: {str(e)}")
            return None
    
    def cleanup_old_files(self, hours_old=1):
        """
        Clean up files older than specified hours
        
        Args:
            hours_old: Number of hours after which files should be cleaned up
            
        Returns:
            dict: Cleanup results with counts and errors
        """
        try:
            if not self.supabase:
                return {
                    'success': False,
                    'error': 'Storage service unavailable',
                    'files_deleted': 0,
                    'errors': []
                }
            
            cutoff_time = datetime.now() - timedelta(hours=hours_old)
            files_to_delete = []
            errors = []
            
            # Get all files in storage
            all_files = self.list_all_files()
            
            for file_info in all_files:
                try:
                    # Check if file is old enough to delete
                    file_created = file_info.get('created_at')
                    if file_created:
                        # Parse the timestamp (assuming ISO format)
                        if isinstance(file_created, str):
                            file_date = datetime.fromisoformat(file_created.replace('Z', '+00:00'))
                        else:
                            file_date = file_created
                        
                        # Convert to naive datetime for comparison
                        if file_date.tzinfo:
                            file_date = file_date.replace(tzinfo=None)
                        
                        if file_date < cutoff_time:
                            files_to_delete.append(file_info.get('name'))
                except Exception as e:
                    errors.append(f"Error processing file {file_info.get('name', 'unknown')}: {str(e)}")
            
            # Delete old files
            deleted_count = 0
            if files_to_delete:
                try:
                    delete_response = self.supabase.storage.from_(self.bucket_name).remove(files_to_delete)
                    if delete_response.status_code == 200:
                        deleted_count = len(files_to_delete)
                        logger.info(f"Successfully deleted {deleted_count} old files")
                    else:
                        errors.append(f"Failed to delete files: {delete_response}")
                except Exception as e:
                    errors.append(f"Error deleting files: {str(e)}")
            
            return {
                'success': len(errors) == 0,
                'files_deleted': deleted_count,
                'files_checked': len(all_files),
                'errors': errors,
                'cutoff_time': cutoff_time.isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error during file cleanup: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'files_deleted': 0,
                'errors': [str(e)]
            }
    
    def cleanup_session_files_advanced(self, session_key, include_outputs=True):
        """
        Advanced cleanup for session files with detailed reporting
        
        Args:
            session_key: Session key to clean up
            include_outputs: Whether to include generated output files
            
        Returns:
            dict: Detailed cleanup results
        """
        try:
            if not self.supabase:
                return {
                    'success': False,
                    'error': 'Storage service unavailable',
                    'files_deleted': 0
                }
            
            # List all files in the session folder
            session_files = self.list_all_files(session_key)
            files_to_delete = []
            
            for file_info in session_files:
                file_name = file_info.get('name', '')
                file_path = f"{session_key}/{file_name}"
                
                # Include all files or filter out outputs based on parameter
                if include_outputs or not self._is_output_file(file_name):
                    files_to_delete.append(file_path)
            
            # Delete files
            deleted_count = 0
            errors = []
            
            if files_to_delete:
                try:
                    delete_response = self.supabase.storage.from_(self.bucket_name).remove(files_to_delete)
                    if delete_response.status_code == 200:
                        deleted_count = len(files_to_delete)
                        logger.info(f"Successfully deleted {deleted_count} files for session {session_key}")
                    else:
                        errors.append(f"Failed to delete session files: {delete_response}")
                except Exception as e:
                    errors.append(f"Error deleting session files: {str(e)}")
            
            return {
                'success': len(errors) == 0,
                'session_key': session_key,
                'files_deleted': deleted_count,
                'files_found': len(session_files),
                'errors': errors
            }
            
        except Exception as e:
            logger.error(f"Error cleaning up session {session_key}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'files_deleted': 0
            }
    
    def get_storage_stats(self):
        """
        Get storage usage statistics
        
        Returns:
            dict: Storage statistics
        """
        try:
            if not self.supabase:
                return {
                    'success': False,
                    'error': 'Storage service unavailable'
                }
            
            all_files = self.list_all_files()
            total_files = len(all_files)
            total_size = 0
            file_types = {}
            old_files_count = 0
            
            cutoff_time = datetime.now() - timedelta(hours=1)
            
            for file_info in all_files:
                # Count file size
                file_size = file_info.get('metadata', {}).get('size', 0)
                total_size += file_size
                
                # Count file types
                file_name = file_info.get('name', '')
                file_ext = os.path.splitext(file_name)[1].lower()
                file_types[file_ext] = file_types.get(file_ext, 0) + 1
                
                # Count old files
                try:
                    file_created = file_info.get('created_at')
                    if file_created:
                        if isinstance(file_created, str):
                            file_date = datetime.fromisoformat(file_created.replace('Z', '+00:00'))
                        else:
                            file_date = file_created
                        
                        if file_date.tzinfo:
                            file_date = file_date.replace(tzinfo=None)
                        
                        if file_date < cutoff_time:
                            old_files_count += 1
                except:
                    pass
            
            return {
                'success': True,
                'total_files': total_files,
                'total_size_bytes': total_size,
                'total_size_mb': round(total_size / (1024 * 1024), 2),
                'file_types': file_types,
                'old_files_count': old_files_count,
                'cleanup_recommended': old_files_count > 0
            }
            
        except Exception as e:
            logger.error(f"Error getting storage stats: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _is_output_file(self, filename):
        """
        Check if a file is a generated output file
        
        Args:
            filename: Name of the file
            
        Returns:
            bool: True if it's an output file
        """
        output_extensions = ['.xlsx', '.pdf', '.docx']
        file_ext = os.path.splitext(filename)[1].lower()
        return file_ext in output_extensions


class FileCleanupService:
    """Service for managing file cleanup and storage maintenance"""
    
    def __init__(self):
        self.storage_service = SupabaseStorageService()
    
    def cleanup_expired_files(self, hours_old=1):
        """
        Clean up files older than specified hours and update database records
        
        Args:
            hours_old: Number of hours after which files should be cleaned up
            
        Returns:
            dict: Cleanup results with detailed information
        """
        from .models import ProcessedDocument, UserSession
        
        try:
            # Clean up storage files (handle storage service unavailability gracefully)
            storage_cleanup = self.storage_service.cleanup_old_files(hours_old)
            storage_success = storage_cleanup.get('success', False)
            
            # Always try to clean up database records even if storage fails
            cutoff_time = datetime.now() - timedelta(hours=hours_old)
            
            # Find old sessions
            old_sessions = UserSession.objects.filter(
                last_activity__lt=cutoff_time
            )
            
            # Clean up associated documents
            old_documents = ProcessedDocument.objects.filter(
                session__in=old_sessions
            )
            
            # Count before deletion
            sessions_count = old_sessions.count()
            documents_count = old_documents.count()
            
            # Delete old documents and sessions
            old_documents.delete()
            old_sessions.delete()
            
            logger.info(f"Cleaned up {sessions_count} old sessions and {documents_count} documents")
            
            # Consider cleanup successful if database cleanup worked, even if storage failed
            overall_success = True
            if not storage_success and storage_cleanup.get('error') != 'Storage service unavailable':
                overall_success = False
            
            return {
                'success': overall_success,
                'storage_cleanup': storage_cleanup,
                'database_cleanup': {
                    'success': True,
                    'sessions_deleted': sessions_count,
                    'documents_deleted': documents_count
                },
                'total_files_deleted': storage_cleanup.get('files_deleted', 0),
                'cutoff_time': cutoff_time.isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error during expired files cleanup: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'storage_cleanup': {},
                'database_cleanup': {
                    'sessions_deleted': 0,
                    'documents_deleted': 0
                }
            }
    
    def cleanup_session_manually(self, session_key, cleanup_outputs=True):
        """
        Manually clean up a specific session and its files
        
        Args:
            session_key: Session key to clean up
            cleanup_outputs: Whether to clean up generated output files
            
        Returns:
            dict: Cleanup results
        """
        from .models import ProcessedDocument, UserSession
        
        try:
            # Clean up storage files for the session (handle storage unavailability gracefully)
            storage_cleanup = self.storage_service.cleanup_session_files_advanced(
                session_key, 
                include_outputs=cleanup_outputs
            )
            storage_success = storage_cleanup.get('success', False)
            
            # Always try to clean up database records
            try:
                session = UserSession.objects.get(session_key=session_key)
                documents_count = session.documents.count()
                
                # Delete documents and session
                session.documents.all().delete()
                session.delete()
                
                database_success = True
                database_error = None
            except UserSession.DoesNotExist:
                documents_count = 0
                database_success = True
                database_error = "Session not found in database"
            except Exception as e:
                documents_count = 0
                database_success = False
                database_error = str(e)
            
            # Consider cleanup successful if database cleanup worked, even if storage failed
            overall_success = database_success
            if not storage_success and storage_cleanup.get('error') != 'Storage service unavailable':
                overall_success = False
            
            return {
                'success': overall_success,
                'session_key': session_key,
                'storage_cleanup': storage_cleanup,
                'database_cleanup': {
                    'success': database_success,
                    'documents_deleted': documents_count,
                    'error': database_error
                }
            }
            
        except Exception as e:
            logger.error(f"Error during manual session cleanup: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'session_key': session_key
            }
    
    def get_cleanup_candidates(self, hours_old=1):
        """
        Get list of sessions and files that are candidates for cleanup
        
        Args:
            hours_old: Number of hours to consider for cleanup
            
        Returns:
            dict: Information about cleanup candidates
        """
        from .models import ProcessedDocument, UserSession
        
        try:
            cutoff_time = datetime.now() - timedelta(hours=hours_old)
            
            # Find old sessions
            old_sessions = UserSession.objects.filter(
                last_activity__lt=cutoff_time
            ).values('session_key', 'last_activity', 'is_active')
            
            # Get storage statistics
            storage_stats = self.storage_service.get_storage_stats()
            
            # Count documents for old sessions
            old_documents_count = ProcessedDocument.objects.filter(
                session__last_activity__lt=cutoff_time
            ).count()
            
            return {
                'success': True,
                'cutoff_time': cutoff_time.isoformat(),
                'old_sessions': list(old_sessions),
                'old_sessions_count': len(old_sessions),
                'old_documents_count': old_documents_count,
                'storage_stats': storage_stats,
                'cleanup_recommended': len(old_sessions) > 0 or storage_stats.get('old_files_count', 0) > 0
            }
            
        except Exception as e:
            logger.error(f"Error getting cleanup candidates: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def schedule_automatic_cleanup(self):
        """
        Perform automatic cleanup as part of scheduled maintenance
        This method is designed to be called by management commands or cron jobs
        
        Returns:
            dict: Cleanup results
        """
        logger.info("Starting scheduled automatic cleanup")
        
        # Perform cleanup of files older than 1 hour
        cleanup_result = self.cleanup_expired_files(hours_old=1)
        
        # Log results
        if cleanup_result.get('success'):
            files_deleted = cleanup_result.get('total_files_deleted', 0)
            sessions_deleted = cleanup_result.get('database_cleanup', {}).get('sessions_deleted', 0)
            logger.info(f"Automatic cleanup completed: {files_deleted} files, {sessions_deleted} sessions deleted")
        else:
            logger.error(f"Automatic cleanup failed: {cleanup_result.get('error', 'Unknown error')}")
        
        return cleanup_result


class SessionService:
    """Service for managing user sessions and concurrent limits"""
    
    @staticmethod
    def get_or_create_session(request):
        """
        Get or create user session with concurrent limit check
        
        Args:
            request: Django request object
            
        Returns:
            tuple: (UserSession object, created boolean, error message)
        """
        from .models import UserSession
        
        session_key = request.session.session_key
        if not session_key:
            request.session.create()
            session_key = request.session.session_key
        
        try:
            # Try to get existing session
            user_session = UserSession.objects.get(session_key=session_key)
            if not user_session.is_active:
                user_session.is_active = True
                user_session.save()
            return user_session, False, None
            
        except UserSession.DoesNotExist:
            # Check concurrent user limit
            active_count = UserSession.get_active_session_count()
            if active_count >= 4:
                return None, False, "System is at capacity (4 users). Please try again later."
            
            # Create new session
            user_session = UserSession.objects.create(
                session_key=session_key,
                is_active=True
            )
            return user_session, True, None
    
    @staticmethod
    def cleanup_inactive_sessions():
        """Clean up sessions inactive for more than 1 hour"""
        from .models import UserSession
        
        cutoff_time = datetime.now() - timedelta(hours=1)
        inactive_sessions = UserSession.objects.filter(
            last_activity__lt=cutoff_time,
            is_active=True
        )
        
        # Clean up storage for inactive sessions using enhanced cleanup
        cleanup_service = FileCleanupService()
        for session in inactive_sessions:
            cleanup_service.cleanup_session_manually(session.session_key, cleanup_outputs=True)
        
        # Deactivate sessions (this will be handled by cleanup_session_manually, but kept for safety)
        inactive_sessions.update(is_active=False)


class OCRService:
    """Service for extracting text from images using Tesseract OCR"""
    
    def __init__(self):
        # Configure Tesseract path if needed (Windows)
        if os.name == 'nt':  # Windows
            # Try common Tesseract installation paths
            possible_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                r'C:\Users\{}\AppData\Local\Tesseract-OCR\tesseract.exe'.format(os.getenv('USERNAME', ''))
            ]
            for path in possible_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break
    
    def extract_text_from_image(self, image_file):
        """
        Extract text from image file using OCR with comprehensive error handling
        
        Args:
            image_file: File object (JPG, PNG) or file path
            
        Returns:
            dict: Contains extracted text and confidence info with detailed error handling
        """
        try:
            # Check if Tesseract is available
            if not self._is_tesseract_available():
                return {
                    'success': False,
                    'text': '',
                    'confidence': 0,
                    'error': 'OCR service unavailable',
                    'details': 'Tesseract OCR is not installed or not properly configured',
                    'suggestions': [
                        'Install Tesseract OCR from https://github.com/tesseract-ocr/tesseract',
                        'On Windows: Download installer from GitHub releases',
                        'On Mac: Use "brew install tesseract"',
                        'On Linux: Use "sudo apt-get install tesseract-ocr"'
                    ],
                    'fallback_suggestion': 'You can still process text files (.txt) and text-based PDFs without OCR',
                    'retry_allowed': False
                }
            
            # Load image
            if hasattr(image_file, 'read'):
                # File object
                image_file.seek(0)
                image_data = image_file.read()
                image = Image.open(io.BytesIO(image_data))
            else:
                # File path
                image = Image.open(image_file)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Preprocess image for better OCR accuracy
            processed_image = self._preprocess_image(image)
            
            # Extract text with confidence data
            ocr_data = pytesseract.image_to_data(
                processed_image, 
                output_type=pytesseract.Output.DICT,
                config='--psm 6'  # Assume uniform block of text
            )
            
            # Extract text and calculate confidence
            extracted_text = pytesseract.image_to_string(processed_image, config='--psm 6')
            
            # Calculate average confidence
            confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            # Clean up extracted text
            cleaned_text = self._clean_extracted_text(extracted_text)
            
            # Check if we got meaningful text
            if not cleaned_text.strip():
                return {
                    'success': False,
                    'text': '',
                    'confidence': 0,
                    'error': 'No readable text found',
                    'details': 'OCR could not detect any readable text in the image',
                    'suggestions': [
                        'Ensure the image is clear and well-lit',
                        'Check that text is large enough and not blurry',
                        'Try adjusting image brightness/contrast',
                        'Ensure text is in a supported language (English)',
                        'Consider scanning at higher resolution'
                    ],
                    'retry_allowed': True,
                    'fallback_suggestion': 'If the document contains text, try converting it to PDF or typing the content into a text file'
                }
            
            return {
                'success': True,
                'text': cleaned_text,
                'confidence': avg_confidence,
                'word_count': len(cleaned_text.split()),
                'message': f'Text extracted successfully with {avg_confidence:.1f}% confidence'
            }
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            
            # Provide specific error messages for common issues
            error_msg = str(e).lower()
            if 'tesseract' in error_msg:
                return {
                    'success': False,
                    'text': '',
                    'confidence': 0,
                    'error': 'OCR service error',
                    'details': 'Tesseract OCR encountered an error during processing',
                    'suggestions': [
                        'Try uploading a different image format (JPG, PNG)',
                        'Ensure the image is not corrupted',
                        'Check that the image contains readable text'
                    ],
                    'retry_allowed': True,
                    'installation_help': 'If Tesseract is not installed: https://github.com/tesseract-ocr/tesseract'
                }
            elif 'image' in error_msg or 'pil' in error_msg or 'cannot identify image file' in error_msg:
                return {
                    'success': False,
                    'text': '',
                    'confidence': 0,
                    'error': 'Invalid image file',
                    'details': 'The uploaded file is not a valid image or is corrupted',
                    'suggestions': [
                        'Check that the file is a valid JPG, PNG, or PDF',
                        'Try opening the file on your computer to verify it works',
                        'Re-save or re-export the image from your source application',
                        'Try uploading a different file'
                    ],
                    'retry_allowed': True
                }
            elif 'memory' in error_msg or 'size' in error_msg:
                return {
                    'success': False,
                    'text': '',
                    'confidence': 0,
                    'error': 'Image too large to process',
                    'details': 'The image is too large for OCR processing',
                    'suggestions': [
                        'Resize the image to a smaller resolution',
                        'Compress the image file',
                        'Split large documents into smaller sections'
                    ],
                    'retry_allowed': True
                }
            else:
                return {
                    'success': False,
                    'text': '',
                    'confidence': 0,
                    'error': 'OCR processing failed',
                    'details': f'Unexpected error during text extraction: {str(e)}',
                    'suggestions': [
                        'Try uploading a different image',
                        'Check that the image is clear and readable',
                        'Try again in a few moments'
                    ],
                    'retry_allowed': True
                }
    
    def extract_text_from_pdf(self, pdf_file):
        """
        Extract text from PDF file, using OCR for image-based PDFs
        
        Args:
            pdf_file: PDF file object or file path
            
        Returns:
            dict: Contains extracted text and processing info
        """
        try:
            # Load PDF
            if hasattr(pdf_file, 'read'):
                pdf_file.seek(0)
                pdf_data = pdf_file.read()
                pdf_document = fitz.open(stream=pdf_data, filetype="pdf")
            else:
                pdf_document = fitz.open(pdf_file)
            
            all_text = []
            total_confidence = 0
            page_count = 0
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Try to extract text directly first
                page_text = page.get_text()
                
                if page_text.strip():
                    # Text-based PDF
                    all_text.append(page_text)
                    total_confidence += 95  # High confidence for text-based PDFs
                else:
                    # Image-based PDF - use OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                    
                    ocr_result = self.extract_text_from_image(image)
                    if ocr_result['success']:
                        all_text.append(ocr_result['text'])
                        total_confidence += ocr_result['confidence']
                    else:
                        logger.warning(f"OCR failed for PDF page {page_num + 1}")
                
                page_count += 1
            
            pdf_document.close()
            
            combined_text = '\n\n'.join(all_text)
            avg_confidence = total_confidence / page_count if page_count > 0 else 0
            
            return {
                'success': True,
                'text': combined_text,
                'confidence': avg_confidence,
                'page_count': page_count,
                'word_count': len(combined_text.split()),
                'message': f'PDF processed successfully ({page_count} pages, {avg_confidence:.1f}% confidence)'
            }
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {str(e)}")
            return {
                'success': False,
                'text': '',
                'confidence': 0,
                'error': 'PDF processing failed',
                'details': str(e)
            }
    
    def _preprocess_image(self, image):
        """
        Preprocess image to improve OCR accuracy
        
        Args:
            image: PIL Image object
            
        Returns:
            PIL Image: Preprocessed image
        """
        try:
            # Resize image if too small (OCR works better on larger images)
            width, height = image.size
            if width < 1000 or height < 1000:
                scale_factor = max(1000 / width, 1000 / height)
                new_width = int(width * scale_factor)
                new_height = int(height * scale_factor)
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Convert to grayscale
            image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            
            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2.0)
            
            # Apply slight blur to reduce noise
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            return image
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed, using original: {str(e)}")
            return image
    
    def _clean_extracted_text(self, text):
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw OCR text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(lines)
        
        # Remove multiple spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def _is_tesseract_available(self):
        """
        Check if Tesseract OCR is available on the system
        
        Returns:
            bool: True if Tesseract is available
        """
        try:
            # Try to get Tesseract version
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False
    
    def process_file(self, file_obj, file_type):
        """
        Process file based on type and extract text
        
        Args:
            file_obj: File object to process
            file_type: File extension (jpg, png, pdf, txt)
            
        Returns:
            dict: Processing result with extracted text
        """
        file_type = file_type.lower()
        
        try:
            if file_type in ['jpg', 'jpeg', 'png']:
                return self.extract_text_from_image(file_obj)
            elif file_type == 'pdf':
                return self.extract_text_from_pdf(file_obj)
            elif file_type == 'txt':
                # For text files, just read content
                file_obj.seek(0)
                content = file_obj.read()
                if isinstance(content, bytes):
                    content = content.decode('utf-8', errors='ignore')
                
                return {
                    'success': True,
                    'text': content,
                    'confidence': 100,  # Text files have perfect "confidence"
                    'word_count': len(content.split()),
                    'message': 'Text file processed successfully'
                }
            else:
                return {
                    'success': False,
                    'text': '',
                    'confidence': 0,
                    'error': f'Unsupported file type: {file_type}',
                    'details': 'Supported types: JPG, PNG, PDF, TXT'
                }
                
        except Exception as e:
            logger.error(f"File processing failed for {file_type}: {str(e)}")
            return {
                'success': False,
                'text': '',
                'confidence': 0,
                'error': 'File processing failed',
                'details': str(e)
            }


class LLMService:
    """Service for parsing extracted text using LLM APIs (OpenAI and Gemini)"""
    
    def __init__(self):
        # Initialize OpenAI client
        self.openai_client = None
        if hasattr(settings, 'OPENAI_API_KEY') and settings.OPENAI_API_KEY:
            self.openai_client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        
        # Initialize Gemini client
        self.gemini_client = None
        if hasattr(settings, 'GEMINI_API_KEY') and settings.GEMINI_API_KEY:
            genai.configure(api_key=settings.GEMINI_API_KEY)
            self.gemini_client = genai.GenerativeModel('gemini-1.5-flash')
    
    def parse_banking_document(self, text: str, document_type: str = "banking_document") -> Dict[str, Any]:
        """
        Parse banking document text using LLM to extract structured data
        
        Args:
            text: Extracted text from document
            document_type: Type of document (banking_document, loan_application, etc.)
            
        Returns:
            dict: Structured banking data with confidence scores
        """
        if not text or not text.strip():
            return {
                'success': False,
                'error': 'No text provided for parsing',
                'data': {}
            }
        
        # Try Gemini first, then OpenAI as fallback
        result = self._parse_with_gemini(text, document_type)
        if not result['success'] and self.openai_client:
            logger.info("Gemini parsing failed, trying OpenAI")
            result = self._parse_with_openai(text, document_type)
        
        return result
    
    def _parse_with_gemini(self, text: str, document_type: str) -> Dict[str, Any]:
        """Parse text using Gemini API with comprehensive error handling"""
        if not self.gemini_client:
            return {
                'success': False,
                'error': 'AI service unavailable',
                'details': 'Gemini API is not configured or API key is missing',
                'suggestions': [
                    'Check that GEMINI_API_KEY is set in environment variables',
                    'Verify API key is valid and has sufficient quota',
                    'Contact administrator if problem persists'
                ],
                'retry_allowed': False,
                'data': {}
            }
        
        try:
            prompt = self._build_parsing_prompt(text, document_type)
            
            # Generate response with Gemini
            response = self.gemini_client.generate_content(prompt)
            
            if not response.text:
                return {
                    'success': False,
                    'error': 'Empty response from Gemini API',
                    'data': {}
                }
            
            # Parse JSON response
            parsed_data = self._parse_llm_response(response.text)
            
            if parsed_data:
                return {
                    'success': True,
                    'data': parsed_data,
                    'provider': 'gemini',
                    'confidence': parsed_data.get('confidence_score', 0.8),
                    'message': 'Document parsed successfully with Gemini'
                }
            else:
                return {
                    'success': False,
                    'error': 'Failed to parse Gemini response as JSON',
                    'data': {},
                    'raw_response': response.text[:500]  # First 500 chars for debugging
                }
                
        except Exception as e:
            logger.error(f"Gemini API error: {str(e)}")
            error_msg = str(e).lower()
            
            # Provide specific error messages for common API issues
            if 'quota' in error_msg or 'limit' in error_msg:
                return {
                    'success': False,
                    'error': 'AI service quota exceeded',
                    'details': 'The AI service has reached its usage limit',
                    'suggestions': [
                        'Try again in a few minutes',
                        'Contact administrator about increasing quota',
                        'Try processing a smaller document'
                    ],
                    'retry_allowed': True,
                    'data': {}
                }
            elif 'api key' in error_msg or 'authentication' in error_msg or 'unauthorized' in error_msg:
                return {
                    'success': False,
                    'error': 'AI service authentication failed',
                    'details': 'Invalid or expired API key',
                    'suggestions': [
                        'Contact administrator to check API key configuration',
                        'Verify API key has proper permissions'
                    ],
                    'retry_allowed': False,
                    'data': {}
                }
            elif 'network' in error_msg or 'connection' in error_msg or 'timeout' in error_msg:
                return {
                    'success': False,
                    'error': 'AI service connection failed',
                    'details': 'Could not connect to AI service',
                    'suggestions': [
                        'Check your internet connection',
                        'Try again in a few minutes',
                        'Contact support if problem persists'
                    ],
                    'retry_allowed': True,
                    'data': {}
                }
            else:
                return {
                    'success': False,
                    'error': 'AI service error',
                    'details': f'Unexpected error from AI service: {str(e)}',
                    'suggestions': [
                        'Try again in a few minutes',
                        'Try processing a different document',
                        'Contact support if problem persists'
                    ],
                    'retry_allowed': True,
                    'data': {}
                }
    
    def _parse_with_openai(self, text: str, document_type: str) -> Dict[str, Any]:
        """Parse text using OpenAI API"""
        if not self.openai_client:
            return {
                'success': False,
                'error': 'OpenAI API not configured',
                'data': {}
            }
        
        try:
            prompt = self._build_parsing_prompt(text, document_type)
            
            # Generate response with OpenAI
            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a banking document analysis expert. Extract structured data from documents and return valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.1  # Low temperature for consistent output
            )
            
            response_text = response.choices[0].message.content
            
            if not response_text:
                return {
                    'success': False,
                    'error': 'Empty response from OpenAI API',
                    'data': {}
                }
            
            # Parse JSON response
            parsed_data = self._parse_llm_response(response_text)
            
            if parsed_data:
                return {
                    'success': True,
                    'data': parsed_data,
                    'provider': 'openai',
                    'confidence': parsed_data.get('confidence_score', 0.8),
                    'message': 'Document parsed successfully with OpenAI'
                }
            else:
                return {
                    'success': False,
                    'error': 'Failed to parse OpenAI response as JSON',
                    'data': {},
                    'raw_response': response_text[:500]  # First 500 chars for debugging
                }
                
        except Exception as e:
            logger.error(f"OpenAI API error: {str(e)}")
            error_msg = str(e).lower()
            
            # Provide specific error messages for common OpenAI API issues
            if 'quota' in error_msg or 'billing' in error_msg or 'insufficient_quota' in error_msg:
                return {
                    'success': False,
                    'error': 'AI service quota exceeded',
                    'details': 'OpenAI API quota has been exceeded',
                    'suggestions': [
                        'Check OpenAI account billing and usage',
                        'Try again next month or add credits',
                        'Contact administrator about API limits'
                    ],
                    'retry_allowed': False,
                    'data': {}
                }
            elif 'api key' in error_msg or 'authentication' in error_msg or 'unauthorized' in error_msg:
                return {
                    'success': False,
                    'error': 'AI service authentication failed',
                    'details': 'OpenAI API key is invalid or expired',
                    'suggestions': [
                        'Contact administrator to check API key',
                        'Verify API key has proper permissions'
                    ],
                    'retry_allowed': False,
                    'data': {}
                }
            elif 'rate limit' in error_msg or 'too many requests' in error_msg:
                return {
                    'success': False,
                    'error': 'AI service rate limit exceeded',
                    'details': 'Too many requests sent to OpenAI API',
                    'suggestions': [
                        'Wait a few minutes before trying again',
                        'Try processing fewer documents simultaneously'
                    ],
                    'retry_allowed': True,
                    'data': {}
                }
            elif 'network' in error_msg or 'connection' in error_msg or 'timeout' in error_msg:
                return {
                    'success': False,
                    'error': 'AI service connection failed',
                    'details': 'Could not connect to OpenAI API',
                    'suggestions': [
                        'Check your internet connection',
                        'Try again in a few minutes',
                        'Contact support if problem persists'
                    ],
                    'retry_allowed': True,
                    'data': {}
                }
            else:
                return {
                    'success': False,
                    'error': 'AI service error',
                    'details': f'Unexpected error from OpenAI API: {str(e)}',
                    'suggestions': [
                        'Try again in a few minutes',
                        'Try processing a different document',
                        'Contact support if problem persists'
                    ],
                    'retry_allowed': True,
                    'data': {}
                }
    
    def _build_parsing_prompt(self, text: str, document_type: str) -> str:
        """Build structured prompt for LLM parsing"""
        
        base_prompt = f"""
Analyze the following {document_type} text and extract structured banking information. 
Return ONLY valid JSON with the following structure:

{{
    "document_type": "detected document type",
    "confidence_score": 0.0-1.0,
    "personal_information": {{
        "full_name": "extracted name or null",
        "account_number": "account number or null",
        "customer_id": "customer ID or null",
        "address": "full address or null",
        "phone": "phone number or null",
        "email": "email address or null"
    }},
    "financial_data": {{
        "account_balance": "current balance or null",
        "available_balance": "available balance or null",
        "transactions": [
            {{
                "date": "transaction date",
                "description": "transaction description",
                "amount": "amount with sign",
                "type": "debit/credit/transfer"
            }}
        ],
        "monthly_summary": {{
            "total_deposits": "total deposits or null",
            "total_withdrawals": "total withdrawals or null",
            "fees_charged": "fees charged or null"
        }}
    }},
    "loan_information": {{
        "loan_amount": "loan amount or null",
        "interest_rate": "interest rate or null",
        "loan_term": "loan term or null",
        "monthly_payment": "monthly payment or null",
        "remaining_balance": "remaining balance or null"
    }},
    "dates": {{
        "statement_date": "statement date or null",
        "statement_period": "statement period or null",
        "due_date": "payment due date or null"
    }},
    "bank_information": {{
        "bank_name": "bank name or null",
        "branch": "branch information or null",
        "routing_number": "routing number or null",
        "swift_code": "SWIFT code or null"
    }},
    "extracted_text_quality": {{
        "clarity": "high/medium/low",
        "completeness": "complete/partial/incomplete",
        "issues": ["list of any issues found"]
    }}
}}

Important instructions:
1. Return ONLY the JSON object, no additional text
2. Use null for missing information, don't make up data
3. For amounts, include currency symbols if present
4. For dates, use consistent format (YYYY-MM-DD if possible)
5. Set confidence_score based on text clarity and completeness
6. Extract ALL transactions found in the document
7. Be precise with numbers and dates

Document text to analyze:
{text}
"""
        
        return base_prompt.strip()
    
    def _parse_llm_response(self, response_text: str) -> Optional[Dict[str, Any]]:
        """
        Parse LLM response and extract JSON data
        
        Args:
            response_text: Raw response from LLM
            
        Returns:
            dict: Parsed JSON data or None if parsing fails
        """
        try:
            # Clean response text
            cleaned_text = response_text.strip()
            
            # Try to find JSON in the response
            json_start = cleaned_text.find('{')
            json_end = cleaned_text.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_text = cleaned_text[json_start:json_end]
                parsed_data = json.loads(json_text)
                
                # Validate required structure
                if self._validate_parsed_data(parsed_data):
                    return parsed_data
                else:
                    logger.warning("Parsed data failed validation")
                    return None
            else:
                logger.warning("No JSON found in LLM response")
                return None
                
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {str(e)}")
            return None
        except Exception as e:
            logger.error(f"Response parsing error: {str(e)}")
            return None
    
    def _validate_parsed_data(self, data: Dict[str, Any]) -> bool:
        """
        Validate that parsed data has required structure
        
        Args:
            data: Parsed JSON data
            
        Returns:
            bool: True if data structure is valid
        """
        required_keys = [
            'document_type',
            'confidence_score',
            'personal_information',
            'financial_data',
            'dates',
            'bank_information'
        ]
        
        try:
            # Check top-level keys
            for key in required_keys:
                if key not in data:
                    logger.warning(f"Missing required key: {key}")
                    return False
            
            # Check confidence score is valid
            confidence = data.get('confidence_score', 0)
            if not isinstance(confidence, (int, float)) or confidence < 0 or confidence > 1:
                logger.warning(f"Invalid confidence score: {confidence}")
                return False
            
            # Check that nested objects are dictionaries
            nested_objects = ['personal_information', 'financial_data', 'dates', 'bank_information']
            for obj_key in nested_objects:
                if not isinstance(data.get(obj_key), dict):
                    logger.warning(f"Invalid nested object: {obj_key}")
                    return False
            
            return True
            
        except Exception as e:
            logger.error(f"Data validation error: {str(e)}")
            return False
    
    def test_api_connection(self) -> Dict[str, Any]:
        """
        Test LLM API connections
        
        Returns:
            dict: Status of API connections
        """
        results = {
            'gemini': {'available': False, 'error': None},
            'openai': {'available': False, 'error': None}
        }
        
        # Test Gemini
        if self.gemini_client:
            try:
                test_response = self.gemini_client.generate_content("Test connection. Respond with 'OK'.")
                if test_response.text and 'OK' in test_response.text:
                    results['gemini']['available'] = True
                else:
                    results['gemini']['error'] = 'Unexpected response'
            except Exception as e:
                results['gemini']['error'] = str(e)
        else:
            results['gemini']['error'] = 'API key not configured'
        
        # Test OpenAI
        if self.openai_client:
            try:
                test_response = self.openai_client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": "Test connection. Respond with 'OK'."}],
                    max_tokens=10
                )
                if test_response.choices[0].message.content and 'OK' in test_response.choices[0].message.content:
                    results['openai']['available'] = True
                else:
                    results['openai']['error'] = 'Unexpected response'
            except Exception as e:
                results['openai']['error'] = str(e)
        else:
            results['openai']['error'] = 'API key not configured'
        
        return results


class DataStructuringService:
    """Service for organizing and formatting extracted banking data"""
    
    def __init__(self):
        pass
    
    def structure_banking_data(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Structure and format parsed banking data for display and export
        
        Args:
            parsed_data: Raw parsed data from LLM
            
        Returns:
            dict: Structured data with formatting and validation
        """
        try:
            structured_data = {
                'metadata': {
                    'document_type': parsed_data.get('document_type', 'Unknown'),
                    'confidence_score': parsed_data.get('confidence_score', 0.0),
                    'processing_timestamp': datetime.now().isoformat(),
                    'data_quality': self._assess_data_quality(parsed_data)
                },
                'summary': self._create_summary(parsed_data),
                'personal_info': self._format_personal_info(parsed_data.get('personal_information', {})),
                'financial_summary': self._format_financial_summary(parsed_data.get('financial_data', {})),
                'transactions': self._format_transactions(parsed_data.get('financial_data', {}).get('transactions', [])),
                'loan_details': self._format_loan_info(parsed_data.get('loan_information', {})),
                'bank_details': self._format_bank_info(parsed_data.get('bank_information', {})),
                'important_dates': self._format_dates(parsed_data.get('dates', {})),
                'display_tables': self._create_display_tables(parsed_data),
                'validation_results': self._validate_extracted_data(parsed_data)
            }
            
            return {
                'success': True,
                'data': structured_data,
                'message': 'Data structured successfully'
            }
            
        except Exception as e:
            logger.error(f"Data structuring failed: {str(e)}")
            return {
                'success': False,
                'error': 'Data structuring failed',
                'details': str(e),
                'data': {}
            }
    
    def _assess_data_quality(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Assess the quality of extracted data"""
        quality_info = parsed_data.get('extracted_text_quality', {})
        
        # Count non-null fields
        total_fields = 0
        filled_fields = 0
        
        for section in ['personal_information', 'financial_data', 'bank_information', 'dates']:
            section_data = parsed_data.get(section, {})
            if isinstance(section_data, dict):
                for key, value in section_data.items():
                    total_fields += 1
                    if value is not None and str(value).strip():
                        filled_fields += 1
        
        completeness_score = filled_fields / total_fields if total_fields > 0 else 0
        
        return {
            'clarity': quality_info.get('clarity', 'unknown'),
            'completeness': quality_info.get('completeness', 'unknown'),
            'completeness_score': completeness_score,
            'filled_fields': filled_fields,
            'total_fields': total_fields,
            'issues': quality_info.get('issues', [])
        }
    
    def _create_summary(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Create a high-level summary of the document"""
        personal_info = parsed_data.get('personal_information', {})
        financial_data = parsed_data.get('financial_data', {})
        
        return {
            'account_holder': personal_info.get('full_name'),
            'account_number': personal_info.get('account_number'),
            'current_balance': financial_data.get('account_balance'),
            'transaction_count': len(financial_data.get('transactions', [])),
            'document_type': parsed_data.get('document_type'),
            'statement_date': parsed_data.get('dates', {}).get('statement_date')
        }
    
    def _format_personal_info(self, personal_data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Format personal information as key-value pairs"""
        formatted_info = []
        
        field_labels = {
            'full_name': 'Full Name',
            'account_number': 'Account Number',
            'customer_id': 'Customer ID',
            'address': 'Address',
            'phone': 'Phone Number',
            'email': 'Email Address'
        }
        
        for field, label in field_labels.items():
            value = personal_data.get(field)
            if value and str(value).strip():
                formatted_info.append({
                    'field': label,
                    'value': str(value),
                    'type': 'personal'
                })
        
        return formatted_info
    
    def _format_financial_summary(self, financial_data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Format financial summary information"""
        formatted_summary = []
        
        # Account balances
        balance_fields = {
            'account_balance': 'Current Balance',
            'available_balance': 'Available Balance'
        }
        
        for field, label in balance_fields.items():
            value = financial_data.get(field)
            if value and str(value).strip():
                formatted_summary.append({
                    'field': label,
                    'value': str(value),
                    'type': 'balance'
                })
        
        # Monthly summary
        monthly_summary = financial_data.get('monthly_summary', {})
        monthly_fields = {
            'total_deposits': 'Total Deposits',
            'total_withdrawals': 'Total Withdrawals',
            'fees_charged': 'Fees Charged'
        }
        
        for field, label in monthly_fields.items():
            value = monthly_summary.get(field)
            if value and str(value).strip():
                formatted_summary.append({
                    'field': label,
                    'value': str(value),
                    'type': 'monthly_summary'
                })
        
        return formatted_summary
    
    def _format_transactions(self, transactions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Format transaction data with validation"""
        formatted_transactions = []
        
        for i, transaction in enumerate(transactions):
            if not isinstance(transaction, dict):
                continue
            
            formatted_transaction = {
                'id': i + 1,
                'date': transaction.get('date', ''),
                'description': transaction.get('description', ''),
                'amount': transaction.get('amount', ''),
                'type': transaction.get('type', ''),
                'formatted_amount': self._format_amount(transaction.get('amount', '')),
                'is_valid': self._validate_transaction(transaction)
            }
            
            formatted_transactions.append(formatted_transaction)
        
        return formatted_transactions
    
    def _format_loan_info(self, loan_data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Format loan information"""
        formatted_loan = []
        
        loan_fields = {
            'loan_amount': 'Loan Amount',
            'interest_rate': 'Interest Rate',
            'loan_term': 'Loan Term',
            'monthly_payment': 'Monthly Payment',
            'remaining_balance': 'Remaining Balance'
        }
        
        for field, label in loan_fields.items():
            value = loan_data.get(field)
            if value and str(value).strip():
                formatted_loan.append({
                    'field': label,
                    'value': str(value),
                    'type': 'loan'
                })
        
        return formatted_loan
    
    def _format_bank_info(self, bank_data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Format bank information"""
        formatted_bank = []
        
        bank_fields = {
            'bank_name': 'Bank Name',
            'branch': 'Branch',
            'routing_number': 'Routing Number',
            'swift_code': 'SWIFT Code'
        }
        
        for field, label in bank_fields.items():
            value = bank_data.get(field)
            if value and str(value).strip():
                formatted_bank.append({
                    'field': label,
                    'value': str(value),
                    'type': 'bank'
                })
        
        return formatted_bank
    
    def _format_dates(self, dates_data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Format important dates"""
        formatted_dates = []
        
        date_fields = {
            'statement_date': 'Statement Date',
            'statement_period': 'Statement Period',
            'due_date': 'Payment Due Date'
        }
        
        for field, label in date_fields.items():
            value = dates_data.get(field)
            if value and str(value).strip():
                formatted_dates.append({
                    'field': label,
                    'value': str(value),
                    'type': 'date'
                })
        
        return formatted_dates
    
    def _create_display_tables(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Create formatted tables for UI display"""
        return {
            'personal_info_table': self._format_personal_info(parsed_data.get('personal_information', {})),
            'financial_summary_table': self._format_financial_summary(parsed_data.get('financial_data', {})),
            'transactions_table': self._format_transactions(parsed_data.get('financial_data', {}).get('transactions', [])),
            'bank_info_table': self._format_bank_info(parsed_data.get('bank_information', {})),
            'dates_table': self._format_dates(parsed_data.get('dates', {}))
        }
    
    def _format_amount(self, amount_str: str) -> str:
        """Format amount string for display"""
        if not amount_str or not str(amount_str).strip():
            return ''
        
        # Try to extract numeric value and format it
        import re
        amount_clean = re.sub(r'[^\d.-]', '', str(amount_str))
        
        try:
            amount_float = float(amount_clean)
            return f"${amount_float:,.2f}"
        except ValueError:
            return str(amount_str)  # Return original if can't parse
    
    def _validate_transaction(self, transaction: Dict[str, Any]) -> bool:
        """Validate transaction data"""
        required_fields = ['date', 'description', 'amount']
        
        for field in required_fields:
            value = transaction.get(field)
            if not value or not str(value).strip():
                return False
        
        return True
    
    def _validate_extracted_data(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate all extracted data and provide feedback"""
        validation_results = {
            'overall_valid': True,
            'issues': [],
            'warnings': [],
            'field_validation': {}
        }
        
        # Validate personal information
        personal_info = parsed_data.get('personal_information', {})
        if not personal_info.get('full_name'):
            validation_results['warnings'].append('No account holder name found')
        
        if not personal_info.get('account_number'):
            validation_results['warnings'].append('No account number found')
        
        # Validate financial data
        financial_data = parsed_data.get('financial_data', {})
        transactions = financial_data.get('transactions', [])
        
        if not transactions:
            validation_results['warnings'].append('No transactions found')
        else:
            invalid_transactions = [t for t in transactions if not self._validate_transaction(t)]
            if invalid_transactions:
                validation_results['issues'].append(f'{len(invalid_transactions)} invalid transactions found')
        
        # Check confidence score
        confidence = parsed_data.get('confidence_score', 0)
        if confidence < 0.5:
            validation_results['warnings'].append('Low confidence score - data may be inaccurate')
        
        validation_results['overall_valid'] = len(validation_results['issues']) == 0
        
        return validation_results


class FileGenerationService:
    """Service for generating output files in Excel, PDF, and DOC formats"""
    
    def __init__(self):
        self.temp_dir = os.path.join(settings.BASE_DIR, 'temp_files')
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def generate_all_formats(self, structured_data: Dict[str, Any], session_key: str) -> Dict[str, Any]:
        """
        Generate all three output formats (Excel, PDF, DOC) from structured data
        
        Args:
            structured_data: Parsed and structured document data
            session_key: User session key for file organization
            
        Returns:
            dict: Contains file paths and generation status for all formats
        """
        results = {
            'success': True,
            'files': {},
            'errors': []
        }
        
        # Generate unique base filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"banking_document_{session_key}_{timestamp}"
        
        # Generate Excel file
        try:
            excel_generator = ExcelGenerator()
            excel_path = excel_generator.generate_excel(structured_data, base_filename, self.temp_dir)
            results['files']['excel'] = {
                'path': excel_path,
                'filename': os.path.basename(excel_path),
                'size': os.path.getsize(excel_path) if os.path.exists(excel_path) else 0
            }
        except Exception as e:
            logger.error(f"Excel generation failed: {str(e)}")
            results['errors'].append(f"Excel generation failed: {str(e)}")
            results['success'] = False
        
        # Generate PDF file
        try:
            pdf_generator = PDFGenerator()
            pdf_path = pdf_generator.generate_pdf(structured_data, base_filename, self.temp_dir)
            results['files']['pdf'] = {
                'path': pdf_path,
                'filename': os.path.basename(pdf_path),
                'size': os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0
            }
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
            results['errors'].append(f"PDF generation failed: {str(e)}")
            results['success'] = False
        
        # Generate DOC file
        try:
            doc_generator = DOCGenerator()
            doc_path = doc_generator.generate_doc(structured_data, base_filename, self.temp_dir)
            results['files']['doc'] = {
                'path': doc_path,
                'filename': os.path.basename(doc_path),
                'size': os.path.getsize(doc_path) if os.path.exists(doc_path) else 0
            }
        except Exception as e:
            logger.error(f"DOC generation failed: {str(e)}")
            results['errors'].append(f"DOC generation failed: {str(e)}")
            results['success'] = False
        
        return results
    
    def cleanup_temp_files(self, file_paths: List[str]):
        """Clean up temporary files"""
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                logger.warning(f"Failed to cleanup file {file_path}: {str(e)}")


class ExcelGenerator:
    """Generate Excel files with structured banking data"""
    
    def generate_excel(self, data: Dict[str, Any], base_filename: str, output_dir: str) -> str:
        """
        Generate Excel file with structured data and professional formatting
        
        Args:
            data: Structured banking data
            base_filename: Base filename without extension
            output_dir: Output directory path
            
        Returns:
            str: Path to generated Excel file
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Banking Document Data"
        
        # Define styles
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        data_alignment = Alignment(horizontal="left", vertical="center")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        current_row = 1
        
        # Add document header
        ws.merge_cells(f'A{current_row}:D{current_row}')
        ws[f'A{current_row}'] = "Banking Document Analysis Report"
        ws[f'A{current_row}'].font = Font(bold=True, size=16)
        ws[f'A{current_row}'].alignment = Alignment(horizontal="center")
        current_row += 2
        
        # Add generation timestamp
        ws[f'A{current_row}'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws[f'A{current_row}'].font = Font(italic=True)
        current_row += 2
        
        # Personal Information Section
        current_row = self._add_section_to_excel(
            ws, "Personal Information", 
            data.get('personal_information', {}),
            current_row, header_font, header_fill, header_alignment, data_alignment, border
        )
        
        # Financial Data Section
        financial_data = data.get('financial_data', {})
        if financial_data:
            current_row = self._add_section_to_excel(
                ws, "Financial Summary",
                {
                    'Account Balance': financial_data.get('account_balance', 'N/A'),
                    'Available Balance': financial_data.get('available_balance', 'N/A'),
                    'Total Deposits': financial_data.get('monthly_summary', {}).get('total_deposits', 'N/A'),
                    'Total Withdrawals': financial_data.get('monthly_summary', {}).get('total_withdrawals', 'N/A'),
                    'Fees Charged': financial_data.get('monthly_summary', {}).get('fees_charged', 'N/A')
                },
                current_row, header_font, header_fill, header_alignment, data_alignment, border
            )
        
        # Transactions Section
        transactions = financial_data.get('transactions', [])
        if transactions:
            current_row += 1
            
            # Transactions header
            ws[f'A{current_row}'] = "Transaction History"
            ws[f'A{current_row}'].font = Font(bold=True, size=14)
            current_row += 1
            
            # Transaction table headers
            headers = ['Date', 'Description', 'Amount', 'Type']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=current_row, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
            
            current_row += 1
            
            # Transaction data
            for transaction in transactions:
                row_data = [
                    transaction.get('date', ''),
                    transaction.get('description', ''),
                    transaction.get('amount', ''),
                    transaction.get('type', '')
                ]
                
                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=current_row, column=col, value=value)
                    cell.alignment = data_alignment
                    cell.border = border
                
                current_row += 1
        
        # Bank Information Section
        bank_info = data.get('bank_information', {})
        if bank_info:
            current_row = self._add_section_to_excel(
                ws, "Bank Information",
                bank_info,
                current_row, header_font, header_fill, header_alignment, data_alignment, border
            )
        
        # Dates Section
        dates_info = data.get('dates', {})
        if dates_info:
            current_row = self._add_section_to_excel(
                ws, "Important Dates",
                dates_info,
                current_row, header_font, header_fill, header_alignment, data_alignment, border
            )
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Save file
        output_path = os.path.join(output_dir, f"{base_filename}.xlsx")
        wb.save(output_path)
        
        return output_path
    
    def _add_section_to_excel(self, ws, section_title: str, section_data: Dict[str, Any], 
                             start_row: int, header_font, header_fill, header_alignment, 
                             data_alignment, border) -> int:
        """Add a data section to Excel worksheet"""
        current_row = start_row + 1
        
        # Section title
        ws[f'A{current_row}'] = section_title
        ws[f'A{current_row}'].font = Font(bold=True, size=14)
        current_row += 1
        
        # Section headers
        ws[f'A{current_row}'] = "Field"
        ws[f'B{current_row}'] = "Value"
        
        for col in ['A', 'B']:
            cell = ws[f'{col}{current_row}']
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        current_row += 1
        
        # Section data
        for key, value in section_data.items():
            if value is not None and str(value).strip():
                # Format field name (convert snake_case to Title Case)
                field_name = key.replace('_', ' ').title()
                
                ws[f'A{current_row}'] = field_name
                ws[f'B{current_row}'] = str(value)
                
                for col in ['A', 'B']:
                    cell = ws[f'{col}{current_row}']
                    cell.alignment = data_alignment
                    cell.border = border
                
                current_row += 1
        
        return current_row


class PDFGenerator:
    """Generate PDF files with structured banking data"""
    
    def generate_pdf(self, data: Dict[str, Any], base_filename: str, output_dir: str) -> str:
        """
        Generate PDF file with professional layout and structured tables
        
        Args:
            data: Structured banking data
            base_filename: Base filename without extension
            output_dir: Output directory path
            
        Returns:
            str: Path to generated PDF file
        """
        output_path = os.path.join(output_dir, f"{base_filename}.pdf")
        
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        # Build PDF content
        story = []
        
        # Title
        story.append(Paragraph("Banking Document Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Generation timestamp
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        story.append(Paragraph(f"Generated: {timestamp}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Personal Information Section
        personal_info = data.get('personal_information', {})
        if personal_info:
            story.append(Paragraph("Personal Information", heading_style))
            table_data = self._dict_to_table_data(personal_info)
            if table_data:
                table = self._create_pdf_table(table_data)
                story.append(table)
                story.append(Spacer(1, 20))
        
        # Financial Summary Section
        financial_data = data.get('financial_data', {})
        if financial_data:
            story.append(Paragraph("Financial Summary", heading_style))
            
            summary_data = {
                'Account Balance': financial_data.get('account_balance', 'N/A'),
                'Available Balance': financial_data.get('available_balance', 'N/A')
            }
            
            monthly_summary = financial_data.get('monthly_summary', {})
            if monthly_summary:
                summary_data.update({
                    'Total Deposits': monthly_summary.get('total_deposits', 'N/A'),
                    'Total Withdrawals': monthly_summary.get('total_withdrawals', 'N/A'),
                    'Fees Charged': monthly_summary.get('fees_charged', 'N/A')
                })
            
            table_data = self._dict_to_table_data(summary_data)
            if table_data:
                table = self._create_pdf_table(table_data)
                story.append(table)
                story.append(Spacer(1, 20))
        
        # Transactions Section
        transactions = financial_data.get('transactions', [])
        if transactions:
            story.append(Paragraph("Transaction History", heading_style))
            
            # Create transaction table
            transaction_data = [['Date', 'Description', 'Amount', 'Type']]
            for transaction in transactions:
                transaction_data.append([
                    transaction.get('date', ''),
                    transaction.get('description', ''),
                    transaction.get('amount', ''),
                    transaction.get('type', '')
                ])
            
            table = Table(transaction_data, colWidths=[1.2*inch, 3*inch, 1*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
            ]))
            
            story.append(table)
            story.append(Spacer(1, 20))
        
        # Bank Information Section
        bank_info = data.get('bank_information', {})
        if bank_info:
            story.append(Paragraph("Bank Information", heading_style))
            table_data = self._dict_to_table_data(bank_info)
            if table_data:
                table = self._create_pdf_table(table_data)
                story.append(table)
                story.append(Spacer(1, 20))
        
        # Important Dates Section
        dates_info = data.get('dates', {})
        if dates_info:
            story.append(Paragraph("Important Dates", heading_style))
            table_data = self._dict_to_table_data(dates_info)
            if table_data:
                table = self._create_pdf_table(table_data)
                story.append(table)
        
        # Build PDF
        doc.build(story)
        
        return output_path
    
    def _dict_to_table_data(self, data_dict: Dict[str, Any]) -> List[List[str]]:
        """Convert dictionary to table data format"""
        table_data = [['Field', 'Value']]
        
        for key, value in data_dict.items():
            if value is not None and str(value).strip():
                field_name = key.replace('_', ' ').title()
                table_data.append([field_name, str(value)])
        
        return table_data if len(table_data) > 1 else []
    
    def _create_pdf_table(self, table_data: List[List[str]]) -> Table:
        """Create a formatted PDF table"""
        table = Table(table_data, colWidths=[2.5*inch, 3.5*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
        ]))
        
        return table


class DOCGenerator:
    """Generate DOC files with structured banking data"""
    
    def generate_doc(self, data: Dict[str, Any], base_filename: str, output_dir: str) -> str:
        """
        Generate DOC file with consistent styling and structured content
        
        Args:
            data: Structured banking data
            base_filename: Base filename without extension
            output_dir: Output directory path
            
        Returns:
            str: Path to generated DOC file
        """
        output_path = os.path.join(output_dir, f"{base_filename}.docx")
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Banking Document Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation timestamp
        timestamp_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_run = timestamp_para.runs[0]
        timestamp_run.italic = True
        
        doc.add_paragraph()  # Add space
        
        # Personal Information Section
        personal_info = data.get('personal_information', {})
        if personal_info:
            doc.add_heading('Personal Information', level=1)
            self._add_dict_as_table(doc, personal_info)
            doc.add_paragraph()
        
        # Financial Summary Section
        financial_data = data.get('financial_data', {})
        if financial_data:
            doc.add_heading('Financial Summary', level=1)
            
            summary_data = {
                'Account Balance': financial_data.get('account_balance', 'N/A'),
                'Available Balance': financial_data.get('available_balance', 'N/A')
            }
            
            monthly_summary = financial_data.get('monthly_summary', {})
            if monthly_summary:
                summary_data.update({
                    'Total Deposits': monthly_summary.get('total_deposits', 'N/A'),
                    'Total Withdrawals': monthly_summary.get('total_withdrawals', 'N/A'),
                    'Fees Charged': monthly_summary.get('fees_charged', 'N/A')
                })
            
            self._add_dict_as_table(doc, summary_data)
            doc.add_paragraph()
        
        # Transactions Section
        transactions = financial_data.get('transactions', [])
        if transactions:
            doc.add_heading('Transaction History', level=1)
            
            # Create transaction table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Date', 'Description', 'Amount', 'Type']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add transaction data
            for transaction in transactions:
                row_cells = table.add_row().cells
                row_cells[0].text = transaction.get('date', '')
                row_cells[1].text = transaction.get('description', '')
                row_cells[2].text = transaction.get('amount', '')
                row_cells[3].text = transaction.get('type', '')
            
            doc.add_paragraph()
        
        # Bank Information Section
        bank_info = data.get('bank_information', {})
        if bank_info:
            doc.add_heading('Bank Information', level=1)
            self._add_dict_as_table(doc, bank_info)
            doc.add_paragraph()
        
        # Important Dates Section
        dates_info = data.get('dates', {})
        if dates_info:
            doc.add_heading('Important Dates', level=1)
            self._add_dict_as_table(doc, dates_info)
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    def _add_dict_as_table(self, doc, data_dict: Dict[str, Any]):
        """Add dictionary data as a formatted table to the document"""
        # Filter out empty values
        filtered_data = {k: v for k, v in data_dict.items() 
                        if v is not None and str(v).strip()}
        
        if not filtered_data:
            doc.add_paragraph("No data available")
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Value'
        
        # Make headers bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for key, value in filtered_data.items():
            row_cells = table.add_row().cells
            field_name = key.replace('_', ' ').title()
            row_cells[0].text = field_name
            row_cells[1].text = str(value)